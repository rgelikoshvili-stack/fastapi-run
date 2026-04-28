from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

def shade_cell(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def add_heading(text, level=1, color='1a1512', size=16):
    p = doc.add_heading('', level=level)
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(int(color[0:2],16),int(color[2:4],16),int(color[4:6],16))
    return p

def h1(t): return add_heading(t, 1, '1a1512', 17)
def h2(t): return add_heading(t, 2, '8c3c2d', 13)
def h3(t): return add_heading(t, 3, '3a2f28', 11)

def para(text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(int(color[0:2],16),int(color[2:4],16),int(color[4:6],16))
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1a,0x15,0x12)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EAF3DE')
    p._p.get_or_add_pPr().append(shd)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.6 + level*0.5)
    return p

def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        r = c.paragraphs[0].runs[0]
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shade_cell(c, '1a1512')
    for row_data in rows:
        row = t.add_row()
        for i, v in enumerate(row_data):
            row.cells[i].text = str(v)
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def PB(): doc.add_page_break()
def SP(): doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BRIDGE HUB — FINANCIAL OS')
r.font.name = 'Calibri'; r.font.size = Pt(30); r.font.bold = True
r.font.color.rgb = RGBColor(0x1a,0x15,0x12)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('სრული ტექნიკური აუდიტი — ყველა ფაილი, კოდი, არქიტექტურა')
r2.font.name = 'Calibri'; r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x8c,0x3c,0x2d)

SP()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(f'{datetime.date.today().strftime("%d %B %Y")}')
r3.font.name = 'Calibri'; r3.font.size = Pt(12); r3.italic = True
r3.font.color.rgb = RGBColor(0x7a,0x6a,0x5a)
SP()

tbl(['პარამეტრი','მნიშვნელობა'],[
    ['Cloud Platform','Google Cloud Run — europe-west1'],
    ['Production URL','https://fastapi-run-226875230147.europe-west1.run.app'],
    ['Backend','FastAPI (Python 3.12)'],
    ['Database','PostgreSQL (Cloud SQL) + psycopg2 pool (min=2, max=10)'],
    ['Frontend','Vanilla JS SPA — 23 static HTML pages, no framework'],
    ['Authentication','JWT Bearer (access + refresh), 5 roles, 20+ permissions'],
    ['DB Security','PostgreSQL Row Level Security (RLS) + tenant_id filter'],
    ['AI/LLM','Anthropic Claude (primary) + Google Gemini (fallback)'],
    ['Active Route Modules','62'],
    ['Service Modules','54'],
    ['DB Migrations','4 (001-004), all in production'],
    ['Integration Tests','40 test files'],
    ['Unit Tests','8 test files'],
],widths=[5,11])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1('1. სისტემის არქიტექტურა')

h2('1.1 ფენების სქემა')
code_block(
'Browser (Vanilla JS SPA — 23 HTML pages)\n'
'  └── HTTPS ──────────────────────────────────────────────\n'
'Google Cloud Run (FastAPI + Uvicorn)\n'
'  ├── [1] add_security_headers   CSP, HSTS, X-Frame-Options\n'
'  ├── [2] tenant_middleware       X-Tenant-ID → request.state.tenant_id\n'
'  ├── [3] auth_middleware         JWT verify → user_id, role\n'
'  ├── [4] rbac_middleware         PERMISSION_MAP → 401/403\n'
'  ├── [5] audit_log_middleware    log to audit_logs table\n'
'  │\n'
'  ├── routes/ (62 modules)        API endpoints\n'
'  │    └── services/ (54 modules) Business logic\n'
'  │         └── db.py             psycopg2 ThreadedConnectionPool\n'
'  │              └── PostgreSQL   RLS via GUC app.current_tenant_id\n'
'  │\n'
'  ├── knowledge/                  Georgian COA + tax rules + JSONL KB\n'
'  └── /static/*.html              Frontend SPA pages\n'
'\n'
'Background loops (asyncio):\n'
'  autopilot_loop()  every 60s    auto-approve high-confidence drafts\n'
'  decay_loop()      every 3600s  reduce pattern confidence over time'
)

h2('1.2 Multi-Tenant Security — 3 ფენა')
tbl(['ფენა','სად','მექანიზმი','გარანტია'],[
    ['Application','routes/*.py services/*.py','tenant_id = getattr(request.state,"tenant_id","default")','Explicit filter in every query'],
    ['Middleware','tenant_middleware.py auth_middleware.py','JWT claim tenant_id + X-Tenant-ID header','Auto-set on every request'],
    ['Database','PostgreSQL RLS','current_setting("app.tenant_id") in WHERE via GUC','Cross-tenant impossible even if app code is buggy'],
],widths=[3,3.5,5,4.5])

h2('1.3 Request Lifecycle (10 ნაბიჯი)')
tbl(['#','ნაბიჯი','ფაილი'],[
    ['1','HTTP Request Cloud Run-ზე შემოდის','main.py'],
    ['2','Security headers (CSP/HSTS/XFrame)','main.py (inline middleware)'],
    ['3','tenant_id resolved (header/param/"default")','middleware/tenant_middleware.py'],
    ['4','JWT Bearer token verify','middleware/auth_middleware.py'],
    ['5','RBAC permission check','middleware/rbac_middleware.py'],
    ['6','Audit log write','middleware/audit_log_middleware.py'],
    ['7','Route handler called','routes/*.py'],
    ['8','Service layer (business logic)','services/*.py'],
    ['9','DB query (GUC set, RLS enforced)','app/api/db.py → PostgreSQL'],
    ['10','JSON response (ok/error pattern)','routes/*.py'],
],widths=[1,7,8])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 2. main.py
# ══════════════════════════════════════════════════════════════════════════════
h1('2. main.py — Entry Point')

h2('2.1 62 Registered Routers')
tbl(['Router Module','Prefix','კატეგორია'],[
    ['routes_health + version_router','GET /health, GET /version','System'],
    ['routes_debug','GET /debug/...','System Debug'],
    ['routes_auth','POST /auth/login, /register, /signup, GET /auth/me','Auth'],
    ['routes_rbac','GET/POST /rbac/...','RBAC Management'],
    ['routes_tenants','GET/POST /tenants/...','Multi-Tenancy'],
    ['routes_bank_csv','POST /bank/upload, GET /bank/history','Bank CSV'],
    ['routes_bank_process','POST /bank/process/...','Bank Processing'],
    ['routes_bank_sync','GET/POST /bank-sync/...','Bank Sync'],
    ['routes_bank_accounts','GET/POST /bank-accounts/...','Bank Accounts'],
    ['routes_approval','GET/POST /approval/...','Approval Queue'],
    ['routes_posting','GET/POST /posting/...','Journal Posting'],
    ['routes_documents','GET/POST /documents/...','Document Upload + Triangle'],
    ['routes_outgoing','GET/POST/PATCH /outgoing/...','Outgoing Invoices'],
    ['routes_ocr','POST /ocr/...','OCR'],
    ['routes_invoice','GET/POST /invoice/...','Invoice (legacy)'],
    ['routes_invoices','GET/POST /invoices/...','CRM Invoices'],
    ['routes_expenses','GET/POST /expenses/...','Expenses'],
    ['routes_contracts','GET/POST /contracts/...','Contracts'],
    ['routes_crm','GET/POST /crm/...','CRM Customers'],
    ['routes_payroll','GET/POST /payroll/...','Payroll'],
    ['routes_budget','GET/POST /budget/...','Budget'],
    ['routes_reports','GET /reports/...','Financial Reports'],
    ['routes_coa','GET /coa/...','Chart of Accounts'],
    ['routes_tax','GET/POST /tax/...','Tax Calculations'],
    ['routes_balance_ge','GET/POST /balance-ge/...','Georgian Balance Sheet'],
    ['routes_1c','GET/POST /1c/...','1C ERP Integration'],
    ['routes_erp_import','POST /erp/import/...','ERP Import'],
    ['routes_erp_connectors','GET/POST /erp/connectors/...','ERP Connectors'],
    ['routes_erp_memory','GET/POST /erp/memory/...','ERP Memory'],
    ['routes_learning','GET/POST /learning/...','AI Learning'],
    ['routes_patterns','GET/POST /patterns/...','AI Patterns'],
    ['routes_decision_engine','GET/POST /decision-engine/...','Decision Engine'],
    ['routes_transaction_ai','GET/POST /transaction-ai/...','Transaction AI'],
    ['routes_ai_chat','GET/POST /api/ai/...','AI Chat & Calculations'],
    ['routes_ai_recommend','GET /api/ai/recommend','AI Recommendations'],
    ['routes_claude_chat','POST /api/claude/...','Claude Direct Chat'],
    ['routes_audit_log','GET /audit-log/...','Audit Log'],
    ['routes_audit','GET /audit/...','Audit'],
    ['routes_audit_engine','GET/POST /audit-engine/...','Audit Engine'],
    ['routes_notifications','GET/POST /notifications/...','Notifications'],
    ['routes_notifications_ws','WS /ws/...','WebSocket Notifications'],
    ['routes_collaboration','GET/POST /collaboration/...','Draft Collaboration'],
    ['routes_dashboard','GET /dashboard/...','Dashboard'],
    ['routes_dashboard_live','GET /dashboard/live/...','Live KPI Dashboard'],
    ['routes_dashboard_insights','GET /dashboard/insights','Dashboard Insights'],
    ['routes_export','GET/POST /export/...','Export'],
    ['routes_export_v2','GET/POST /export/v2/...','Export V2'],
    ['routes_export_journal','GET /export/journal','Journal Export'],
    ['routes_currency','GET/POST /currency/...','Currency Rates + Convert'],
    ['routes_search','GET /search/...','Full-text Search'],
    ['routes_reconciliation','GET/POST /reconciliation/...','Reconciliation'],
    ['routes_webhooks_v2','POST /webhooks/...','Webhooks'],
    ['routes_client_portal','GET/POST /portal/...','Client Portal'],
    ['routes_email_invoice','GET/POST /email-invoice/...','Email Invoice Parser'],
    ['routes_pdf_report','GET /pdf/...','PDF Reports'],
    ['routes_security','GET/POST /security/...','Security Management'],
    ['routes_qa','GET/POST /qa/...','QA Engine'],
    ['routes_transaction_memory','GET/POST /transaction-memory/...','Transaction Memory'],
    ['routes_learning_explain','GET /learning/explain/...','Learning Explainability'],
    ['routes_system','GET/POST /system/...','System Admin'],
    ['routes_ai_journal','GET /ai-journal/...','AI Journal View'],
    ['routes_api_docs','GET /api-docs/...','API Documentation'],
    ['routes_docs','GET /docs-info/...','Docs Info'],
],widths=[5.5,4.5,6])

h2('2.2 Startup Events')
tbl(['Event','Interval','ფუნქცია','მიზანი'],[
    ['Autopilot Loop','60s','autopilot_approve_service("default")','Auto-approve drafts where confidence≥0.80, usage≥5, success_rate≥0.80'],
    ['Decay Loop','3600s','run_decay_service()','Reduce confidence of unused patterns (decay_rate=0.98 per cycle)'],
    ['DB Migrations','Once on start','ALTER TABLE IF NOT EXISTS ...','Idempotent column additions (zero-downtime)'],
    ['JSON→DB Migration','Once on start','migrate_json_to_db()','Migrate learned_rules.json entries → learning_patterns table'],
    ['Knowledge Preload','Once on start','_load_files()','Load JSONL knowledge files into memory for fast lookup'],
],widths=[3.5,2,5,5.5])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 3. MIDDLEWARE
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Middleware Stack')

h2('3.1 Security Headers (main.py inline)')
tbl(['Header','Value','დანიშნულება'],[
    ['X-Content-Type-Options','nosniff','MIME sniffing attack prevention'],
    ['X-Frame-Options','DENY','Clickjacking prevention'],
    ['X-XSS-Protection','1; mode=block','Legacy browser XSS filter'],
    ['Strict-Transport-Security','max-age=31536000; includeSubDomains','HTTPS enforcement (1 year)'],
    ['Content-Security-Policy','default-src self; script-src self unsafe-inline; style-src self unsafe-inline fonts.googleapis.com; font-src self data: fonts.gstatic.com; img-src self data: https:; connect-src self ...','Script/style/font source restriction'],
    ['X-Powered-By','BridgeHub/1.0','Custom server header'],
],widths=[4.5,4,7.5])

h2('3.2 auth_middleware.py')
bullet('Public paths (no JWT needed): /, /docs, /openapi.json, /health/*, /auth/*, /static/*')
bullet('On protected path: extracts Authorization: Bearer {token}')
bullet('Calls verify_token(token, expected_type="access")')
bullet('Sets request.state: authenticated=True, user_id, role, tenant_id')
bullet('On failure: returns 401 {"ok":false, "error":"Unauthorized"}')

h2('3.3 rbac_middleware.py')
bullet('Public prefixes (skip RBAC): /auth/, /docs, /health, /static, /api/ai/, /api/claude/, /dashboard/, /coa/, /debug/ai-routing')
bullet('Checks: authenticated=True → lookup PERMISSION_MAP(method, path_prefix)')
bullet('If permission found: check role_permissions[role]')
bullet('Returns 401 if not authenticated, 403 if role lacks permission')

h2('3.4 audit_log_middleware.py')
bullet('Skips: /docs, /openapi.json, /static/*')
bullet('Captures: action (POST/GET/...), resource (URL path), actor (user_id), role, tenant_id, ip (X-Forwarded-For), HTTP status')
bullet('Writes to: audit_logs table via log_event() — async, non-blocking')

h2('3.5 tenant_middleware.py')
bullet('Priority order: X-Tenant-ID header → ?tenant_id query param → "default"')
bullet('Sets: request.state.tenant_id')
bullet('Also calls: db.set_tenant_guc(tenant_id) → PostgreSQL SET app.current_tenant_id')
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 4. AUTH & RBAC
# ══════════════════════════════════════════════════════════════════════════════
h1('4. Authentication & RBAC')

h2('4.1 Auth Endpoints — routes_auth.py')
tbl(['Method','Path','Rate Limit','Type','Response'],[
    ['POST','/auth/login','5/min','Public','{"access_token", "refresh_token", "token_type":"bearer"}'],
    ['POST','/auth/register','5/min','Public','{"user_id", "tenant_id", "role"}'],
    ['POST','/auth/signup','3/min','Public','{"tenant_id", "slug", "access_token", "refresh_token"}'],
    ['GET','/auth/me','—','JWT Bearer','{"user_id", "email", "role", "tenant_id", "exp"}'],
    ['POST','/auth/refresh','—','Refresh Token','{"access_token", "token_type":"bearer"}'],
],widths=[2,3.5,2.5,3.5,5])

h2('4.2 JWT Token Structure')
code_block(
'# Access Token payload:\n'
'{\n'
'  "sub": "user@example.com",     # user identifier\n'
'  "user_id": "uuid",\n'
'  "tenant_id": "acme-corp",       # tenant slug\n'
'  "role": "accountant",           # RBAC role\n'
'  "type": "access",\n'
'  "exp": 1784000000               # expiry timestamp\n'
'}'
)

h2('4.3 RBAC Roles — app/api/authz.py')
tbl(['Role','Permissions'],[
    ['admin','reports:read/write, posting:read/write, approval:read/write, payroll:read/write, ocr:read/write, notifications:read/write, search:read, tenants:manage, patterns:manage, audit:view, bank:upload/process, export:any, chat:use, dashboard:admin/view'],
    ['accountant','reports:read, posting:read/write, approval:read/write, payroll:read/write, ocr:read/write, notifications:read, search:read, patterns:view, audit:view, bank:upload/process, export:any, chat:use, dashboard:view'],
    ['reviewer','reports:read, posting:read, approval:read/write, ocr:read, notifications:read, search:read, patterns:view, audit:view, export:any, chat:use, dashboard:view'],
    ['viewer','reports:read, posting:read, payroll:read, ocr:read, notifications:read, search:read, patterns:view, audit:view, export:any, dashboard:view'],
    ['ai_supervisor','reports:read, patterns:view/manage, audit:view, dashboard:view, chat:use'],
],widths=[3,13])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 5. DATABASE
# ══════════════════════════════════════════════════════════════════════════════
h1('5. მონაცემთა ბაზა')

h2('5.1 Connection Pool (app/api/db.py)')
tbl(['Config','Value'],[
    ['Type','psycopg2.pool.ThreadedConnectionPool'],
    ['Min connections','2 (always alive)'],
    ['Max connections','10 (configurable via MAX_CONN env)'],
    ['Source','os.environ["DATABASE_URL"]'],
    ['Tenant GUC','SET app.current_tenant_id = %s (on every getconn())'],
    ['GUC Reset','set_config("app.current_tenant_id", "", false) on putconn()'],
    ['Fallback','Direct psycopg2.connect() if pool exhausted'],
    ['Encoding','client_encoding=UTF8'],
],widths=[4,12])

h2('5.2 Migration ისტორია')
tbl(['File','ახალი ცხრილები','ძირითადი ცვლილებები'],[
    ['001_multi_tenant_schema.sql',
     'tenants, counterparties, processed_documents',
     'company_inn/type/vat_payer columns; journal_drafts: +our_role, +operation_type, +counterparty_inn, +journal_entries JSONB; dedup via file_hash'],
    ['002_row_level_security.sql',
     '(RLS policies)',
     'ENABLE ROW LEVEL SECURITY on journal_drafts, bank_transactions, budgets, customers, contracts, expenses, invoices; GUC-based policy USING (tenant_id = current_setting("app.tenant_id",true))'],
    ['003_triangle_schema.sql',
     'waybills, tax_invoices, commercial_invoices, triangle_matches',
     'waybills: series, number, seller/buyer_inn, total, line_items JSONB, match_score; tax_invoices: vat_amount, related_waybill_id; triangle_matches: discrepancies JSONB, match_score 0-100'],
    ['004_outgoing_invoices.sql',
     'outgoing_invoices, invoice_counters',
     'invoice_type (goods|service), buyer_inn, line_items JSONB, generated_waybill_id, generated_tax_invoice_id, journal_entries JSONB; invoice_counters: per-tenant/year auto-numbering'],
],widths=[4,4,8])

h2('5.3 ყველა ცხრილი (30+)')
tbl(['ცხრილი','Tenant?','Migration','ძირითადი სვეტები'],[
    ['tenants','N/A (master)','001','id UUID, slug, company_name, company_inn, subscription_tier, trial_ends_at, is_vat_payer'],
    ['journal_drafts','✓','001+','id, tenant_id, status, amount, account_code, reason, date, counterparty_inn, counterparty_name, journal_entries JSONB, description, source, our_role, operation_type'],
    ['bank_transactions','✓','pre','id, tenant_id, date, amount, description, account_code, fingerprint (dedup hash)'],
    ['learning_patterns','✓','pre','id, tenant_id, pattern_key, account_code, confidence, usage_count, success_count, last_used, source'],
    ['counterparties','✓','001','id, tenant_id, inn, name, type (vendor/customer), is_vat_payer, is_foreign'],
    ['processed_documents','✓','001','id, tenant_id, file_hash, extraction_method, extracted_data JSONB, pipeline_run_id'],
    ['budgets','✓','pre','id, tenant_id, name, year, month, account_code, category, budgeted'],
    ['waybills','✓','003','id, tenant_id, waybill_number, series, seller_inn, buyer_inn, total, status, match_score, line_items JSONB'],
    ['tax_invoices','✓','003','id, tenant_id, invoice_number, series, seller_inn, buyer_inn, vat_amount, status, related_waybill_id'],
    ['commercial_invoices','✓','003','id, tenant_id, invoice_number, seller_inn, buyer_inn, amount, currency'],
    ['triangle_matches','✓','003','id, tenant_id, waybill_id, tax_invoice_id, journal_draft_id, match_score, discrepancies JSONB'],
    ['outgoing_invoices','✓','004','id, tenant_id, invoice_type, buyer_inn, subtotal, vat_amount, status, line_items JSONB, generated_waybill_id, generated_tax_invoice_id, finalized_at'],
    ['invoice_counters','✓','004','tenant_id, year, last_number (serial per tenant/year)'],
    ['customers','✓','CRM','id, tenant_id, name, inn, email, phone, address, created_at'],
    ['contracts','✓','CRM','id, tenant_id, customer_id, title, amount, status, start_date, end_date'],
    ['expenses','✓','CRM','id, tenant_id, category, amount, date, description, account_code'],
    ['invoices (CRM)','✓','CRM','id, tenant_id, customer_id, amount, status, due_date, items JSONB'],
    ['audit_logs','partial','pre','id, tenant_id, action, resource, actor, role, ip, status, created_at'],
    ['pipeline_runs','N','pre','run_id, filename, state, created_at (global table, no RLS)'],
],widths=[4,2,2,8])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 6. APPROVAL SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Approval სისტემა')

h2('6.1 routes_approval.py — Endpoints')
tbl(['Method','Path','Rate','RBAC','აღწერა'],[
    ['GET','/approval/queue','—','approval:read','Status filter + pagination, returns pending_count'],
    ['POST','/approval/approve/{id}','30/min','approval:write','Approve — FOR UPDATE NOWAIT → status=approved'],
    ['POST','/approval/reject/{id}','30/min','approval:write','Reject with reason text → status=rejected'],
    ['POST','/approval/correct/{id}','30/min','approval:write','Correct account_code / journal entries'],
    ['GET','/approval/audit','—','audit:view','Approval history'],
    ['POST','/approval/autopilot','—','approval:write','Trigger autopilot manually'],
    ['POST','/approval/preview','—','approval:read','Preview draft (no state change)'],
],widths=[2,4,2,3,5])

h2('6.2 approval_service.py')
tbl(['ფუნქცია','Input → Output'],[
    ['get_queue_service()','(status, limit, offset, tenant_id) → {items[], total, pending_count}'],
    ['approve_draft_service()','(draft_id, tenant_id) → {ok, draft} | DRAFT_LOCKED 409'],
    ['reject_draft_service()','(draft_id, reason, tenant_id) → {ok, draft} | DRAFT_LOCKED 409'],
    ['get_audit_service()','(limit, offset) → {items[], total}'],
    ['autopilot_approve_service()','(tenant_id) → {approved_count, skipped_count, errors[]}'],
],widths=[5,11])

h2('6.3 Race Condition Protection')
code_block(
'# approval_service.py + correct_draft_service.py\n'
'# Step 1: Pessimistic lock\n'
'SELECT * FROM journal_drafts\n'
'WHERE id = %s AND tenant_id = %s\n'
'FOR UPDATE NOWAIT  -- raises LockNotAvailable immediately if locked\n\n'
'# Step 2: Handle lock failure\n'
'except psycopg2.errors.LockNotAvailable:\n'
'    conn.rollback()\n'
'    return error_response("Draft is being processed", "DRAFT_LOCKED", 409)\n\n'
'# Step 3: HTTP 409 in routes_approval.py\n'
'def _check_locked(result):\n'
'    if isinstance(result, dict) and result.get("code") == "DRAFT_LOCKED":\n'
'        return JSONResponse(status_code=409, content=result)\n'
'return _check_locked(result) or result'
)

h2('6.4 Autopilot Config')
tbl(['Threshold','Value','მიზეზი'],[
    ['MIN_CONFIDENCE','0.80 (80%)','Pattern must be ≥80% confident'],
    ['MIN_USAGE_COUNT','5','Pattern used ≥5 times'],
    ['MIN_SUCCESS_RATE','0.80 (80%)','80%+ human-approved history'],
    ['MAX_PATTERN_AGE','45 days','Patterns older than 45d are excluded'],
    ['SIGNAL_WEIGHTS','approve=1.0, reject=1.5','Rejections count 50% more than approvals'],
],widths=[4.5,3.5,8])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 7. AI & DECISION ENGINE
# ══════════════════════════════════════════════════════════════════════════════
h1('7. AI & Decision Engine')

h2('7.1 AI Pipeline — 3-Step Fallback')
code_block(
'Input: transaction text / document\n'
'  │\n'
'  ├─[1] Exact Rule Match ───────── learning_patterns (exact key match)\n'
'  │       confidence ≥ 0.80?  ──→  journal_drafts INSERT (auto_approved)\n'
'  │\n'
'  ├─[2] Fuzzy Pattern Match ────── pattern_engine.py (Levenshtein + TF-IDF)\n'
'  │       confidence ≥ threshold? → journal_drafts INSERT (drafted)\n'
'  │\n'
'  └─[3] LLM Fallback ──────────── Claude API (primary) / Gemini (fallback)\n'
'          → account_code + confidence\n'
'          → journal_drafts INSERT (drafted, requires human approval)\n'
'          → learning_service.update_pattern() (feedback loop)'
)

h2('7.2 routes_ai_chat.py — Endpoints')
tbl(['Method','Path','Input','Output'],[
    ['GET','/api/ai/stats','—','patterns_count, avg_confidence, auto_rate, total_feedback'],
    ['POST','/api/ai/chat','message, session_id, tenant_id','AI response with context + vector search'],
    ['GET','/api/ai/search','q, top_k=5','Vector search results'],
    ['POST','/api/ai/vat','amount, inclusive (bool), service_type','net, vat, gross, rate'],
    ['POST','/api/ai/dividend','gross_amount','net, withholding_tax (5%), journal_entries'],
    ['POST','/api/ai/payroll','gross, include_employee_payg','net, pit(20%), pension(2%), journal_entries'],
    ['POST','/api/ai/cit','distributed_profit','cit_amount(15%), tax_base, journal_entries'],
    ['POST','/api/ai/depreciation','cost, residual, useful_life_years','annual_depreciation, journal_entries'],
    ['POST','/api/ai/classify','text, amount (opt)','account_code, confidence, pattern_source'],
    ['POST','/api/ai/learn','pattern_key, account_code, confidence','pattern_id, status'],
    ['GET','/api/ai/recommend','—','[{severity, title, description, action}]'],
],widths=[2,4.5,4.5,5])

h2('7.3 Learning System')
tbl(['Component','File','ფუნქცია'],[
    ['Pattern Storage','learning_patterns (DB)','tenant_id, pattern_key, account_code, confidence, usage_count, last_used'],
    ['Feedback Loop','learning_service.py','update_pattern_feedback(pattern_key, signal="approve"|"reject")'],
    ['Confidence Decay','pattern_decay_service.py','Every 3600s: confidence *= 0.98^(days_idle/30). Archive if <0.30'],
    ['Human Feedback','POST /learning/feedback','approve/reject signal → update confidence + success_rate'],
    ['JSON Migration','knowledge_loader.py','migrate_json_to_db() on startup: learned_rules.json → DB'],
],widths=[3.5,4,8.5])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 8. DOCUMENT INTELLIGENCE
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Document Intelligence Pipeline')

h2('8.1 Processing Pipeline')
code_block(
'POST /documents/upload  (PDF / Image / CSV)\n'
'  │\n'
'  ├── [dedup] processed_documents.file_hash → skip if seen\n'
'  ├── [1] ocr_service.py\n'
'  │       native PDF text  ← PyPDF2\n'
'  │       scanned image    ← pytesseract\n'
'  │       low quality      ← Claude Vision API\n'
'  │\n'
'  ├── [2] document_extractor.py\n'
'  │       seller_inn, buyer_inn, date, amount\n'
'  │       series, number, vat_amount\n'
'  │\n'
'  ├── [3] party_resolver.py\n'
'  │       our_role: buyer | seller | foreign\n'
'  │\n'
'  ├── [4] operation_classifier.py\n'
'  │       type: revenue | expense | payroll | transfer | tax | other\n'
'  │\n'
'  ├── [5] doc_journal_builder.py\n'
'  │       journal_entries = [{dr, cr, amount, note}, ...]\n'
'  │\n'
'  └── [6] journal_drafts INSERT\n'
'           status = "drafted" | "auto_approved"\n'
'           processed_documents INSERT (hash dedup)'
)

h2('8.2 Triangle Reconciliation — triangle_matcher.py')
tbl(['Field Compared','Tolerance','Match Weight'],[
    ['seller_inn + buyer_inn','Exact match','40%'],
    ['amount','±5% tolerance','30%'],
    ['date','±7 days tolerance','20%'],
    ['line_items (qty × price)','±2% tolerance','10%'],
],widths=[5,4,7])
code_block(
'# triangle_matcher.py\n'
'match = find_match(waybill_id, tax_invoice_id, journal_draft_id)\n'
'# Returns:\n'
'# { match_score: 0-100,\n'
'#   discrepancies: [{field, waybill_value, tax_inv_value, draft_value}],\n'
'#   status: "matched"|"partial"|"mismatch" }'
)

h2('8.3 Correction Detection — correction_detector.py')
bullet('Pattern: same counterparty_inn + same period (±30d) + similar amount (±20%) + different account_code')
bullet('On detection: original entry gets corrected=True flag')
bullet('UI: correction rows rendered with amber-bg (#FAEEDA) background')
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 9. OUTGOING INVOICES
# ══════════════════════════════════════════════════════════════════════════════
h1('9. Outgoing Invoice სისტემა')

h2('9.1 Lifecycle')
code_block(
'[1] POST /outgoing/drafts\n'
'    → outgoing_invoices INSERT (status=draft)\n'
'\n'
'[2] PATCH /outgoing/drafts/{id}   (2-second debounce auto-save)\n'
'    → UPDATE line_items, buyer_inn, amounts\n'
'\n'
'[3] POST /outgoing/{id}/finalize  (single DB transaction)\n'
'    a. outgoing_invoices UPDATE  status=finalized, finalized_at=NOW()\n'
'    b. invoice_counters UPDATE   last_number++ (per tenant/year)\n'
'    c. waybills INSERT           GEO-{year}-{serial}\n'
'    d. tax_invoices INSERT       TI-{year}-{serial}\n'
'    e. journal_drafts INSERT     status=approved\n'
'         journal_entries JSONB:\n'
'           [{dr:"1210", cr:"6110", amount:subtotal, note:"შემოსავალი"},\n'
'            {dr:"1210", cr:"3310", amount:vat_amount, note:"დღგ 18%"}]'
)

h2('9.2 invoice_creator.py — ძირითადი ფუნქციები')
tbl(['ფუნქცია','Input','Output'],[
    ['create_invoice_draft()','conn, tenant_id, invoice_type, buyer_inn, line_items[]','invoice_id (UUID)'],
    ['update_invoice()','conn, invoice_id, tenant_id, **fields','updated_invoice dict'],
    ['finalize_invoice()','conn, invoice_id, tenant_id','{ invoice_id, waybill_id, tax_invoice_id, journal_entries }'],
    ['list_invoices()','conn, tenant_id, status, limit, offset','{ items[], total }'],
    ['get_invoice()','conn, invoice_id, tenant_id','invoice dict with line_items'],
],widths=[4,5,7])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 10. FINANCIAL REPORTS
# ══════════════════════════════════════════════════════════════════════════════
h1('10. Financial Reports')

h2('10.1 ledger_service.py — 5 Functions')
para('ყველა ფუნქცია კითხულობს journal_drafts.journal_entries JSONB-ს PostgreSQL LATERAL unnest-ით. '
     'მხოლოდ status IN (approved, auto_approved) drafts.')
code_block(
'-- Core SQL pattern used by get_account_ledger():\n'
'SELECT jd.id, jd.date, jd.description, jd.counterparty_name,\n'
'       entry->>\'dr\'  AS dr_account,\n'
'       entry->>\'cr\'  AS cr_account,\n'
'       (entry->>\'amount\')::numeric AS amount,\n'
'       entry->>\'note\' AS note\n'
'FROM journal_drafts jd\n'
'CROSS JOIN LATERAL jsonb_array_elements(jd.journal_entries) AS entry\n'
'WHERE jd.tenant_id = %s\n'
'  AND jd.status IN (\'approved\', \'auto_approved\')\n'
'  AND (entry->>\'dr\' = %s OR entry->>\'cr\' = %s)\n'
'ORDER BY jd.date::date, jd.id'
)
tbl(['ფუნქცია','Input Params','Return'],[
    ['get_account_ledger()','tenant_id, account_code, date_from, date_to','lines[], opening_balance, total_debit, total_credit, closing_balance, count'],
    ['get_trial_balance()','tenant_id, date_from, date_to','accounts[{code,debit,credit,net}], total_debit, total_credit, balanced(bool)'],
    ['get_counterparty_ledger()','tenant_id, counterparty_inn, date_from, date_to','lines[{id,date,desc,amount,journal_entries[]}], total_amount, count'],
    ['get_payroll_ledger()','tenant_id, employee_id(opt), year(opt)','lines[{date,employee,dr,cr,amount,note}], total_wages'],
    ['get_journal_entries()','tenant_id, date(opt), limit, offset','items[], total (paginated)'],
],widths=[4.5,5.5,6])

h2('10.2 All Report Endpoints')
tbl(['Endpoint','Rate','RBAC'],[
    ['GET /reports/monthly','—','reports:read'],
    ['GET /reports/annual','—','reports:read'],
    ['GET /reports/pnl?year=&month=','—','reports:read'],
    ['GET /reports/cashflow?year=&month=','—','reports:read'],
    ['GET /reports/audit-trail','—','reports:read'],
    ['GET /reports/ledger/{account_code}?date_from=&date_to=','10/min','reports:read'],
    ['GET /reports/trial-balance?date_from=&date_to=','10/min','reports:read'],
    ['GET /reports/counterparty/{inn}?date_from=&date_to=','10/min','reports:read'],
    ['GET /reports/payroll?employee_id=&year=','10/min','reports:read'],
    ['GET /reports/journal?date=&limit=&offset=','10/min','reports:read'],
],widths=[8,2,6])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 11. GEORGIAN TAX COMPLIANCE
# ══════════════════════════════════════════════════════════════════════════════
h1('11. Georgian Tax Compliance')

h2('11.1 Tax Rates (app/policy/tax_constants.py)')
tbl(['გადასახადი','განაკვეთი','ბაზა','Journal Entries'],[
    ['VAT / დღგ','18%','Net','Dr 1210 / Cr 6110 (revenue) + Cr 3310 (vat payable)'],
    ['PIT / საშემოსავლო','20%','Gross','Dr 7110 / Cr 3370 (net payable) + Cr 1120 (tax deducted)'],
    ['Employer PAYG / pension','2%','Gross','Dr 7120 / Cr 3380'],
    ['Employee Pension','2%','Gross','Withheld from employee net salary'],
    ['Dividend Withholding','5%','Gross dividend','Dr 3370 / Cr 3350 + Cr 1120'],
    ['Royalty Withholding','10%','Gross','Dr 3XXX / Cr 3350 + Cr 1120'],
    ['Interest Withholding','5%','Gross','Dr 3XXX / Cr 3350'],
    ['CIT (Estonian model)','15%','Distributed profit','Dr 4210 / Cr 3340 → Dr 3340 / Cr 1120'],
    ['Depreciation','cost/life','Annual','Dr 7XXX / Cr 15XX (accumulated depreciation)'],
],widths=[3.5,2.5,3.5,6.5])

h2('11.2 Chart of Accounts Classes')
tbl(['Class','Code Range','კატეგორია'],[
    ['1','1000–1999','Current Assets (cash, receivables, inventory)'],
    ['2','2000–2999','Non-current / Long-term Assets (fixed assets)'],
    ['3','3000–3999','Liabilities (payables, tax payables, loans)'],
    ['4','4000–4999','Equity (share capital, retained earnings)'],
    ['5','5000–5999','Revenue (Georgian: შემოსავლები)'],
    ['6','6000–6999','Other Income'],
    ['7','7000–7999','Expenses (wages, depreciation, operating costs)'],
    ['8','8000–8999','Cost of Goods Sold'],
    ['9','9000–9999','Other income/expenses, write-offs'],
],widths=[2,3,11])

h2('11.3 Key Account Codes')
tbl(['Code','სახელი'],[
    ['1120','Bank account (cash at bank)'],
    ['1210','Accounts receivable (trade debtors)'],
    ['3310','VAT payable'],
    ['3340','Corporate income tax payable'],
    ['3350','Withholding tax payable'],
    ['3370','Salary / dividend payable'],
    ['3380','Pension payable'],
    ['6110','Revenue from services/goods'],
    ['7110','Wages expense'],
    ['7120','Employer pension/social charges'],
],widths=[3,13])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 12. SERVICES FULL INVENTORY
# ══════════════════════════════════════════════════════════════════════════════
h1('12. Services — სრული ინვენტარი (54 მოდული)')

h2('12.1 Core Approval & Posting')
tbl(['ფაილი','ძირითადი ფუნქციები'],[
    ['approval_service.py','get_queue_service, approve_draft_service(FOR UPDATE NOWAIT), reject_draft_service, autopilot_approve_service'],
    ['correct_draft_service.py','correct_draft_service — account/entry correction, SELECT FOR UPDATE NOWAIT, LockNotAvailable→409'],
    ['posting_service.py','post_to_journal — approved draft → ERP/1C API posting'],
    ['posting_preview_service.py','preview_posting — dry-run, shows what would be posted without committing'],
    ['ledger_service.py','get_account_ledger, get_trial_balance, get_counterparty_ledger, get_payroll_ledger, get_journal_entries'],
],widths=[5.5,10.5])

h2('12.2 Document Intelligence')
tbl(['ფაილი','ძირითადი ფუნქციები'],[
    ['ocr_service.py','extract_text_from_pdf (PyPDF2), extract_text_ocr (Tesseract), extract_via_vision_llm (Claude Vision)'],
    ['document_extractor.py','extract_document_data → ExtractedDoc(seller_inn, buyer_inn, date, amount, series, number, vat_amount)'],
    ['party_resolver.py','resolve_our_role(seller_inn, buyer_inn, tenant_inns) → "buyer"|"seller"|"foreign"'],
    ['operation_classifier.py','classify_operation(doc, context) → OperationType: revenue|expense|payroll|transfer|tax|other'],
    ['doc_journal_builder.py','build_journal_from_document(doc, operation_type) → [{dr, cr, amount, note}]'],
    ['triangle_matcher.py','find_match(waybill_id, tax_inv_id, draft_id), calculate_score, detect_discrepancies'],
    ['correction_detector.py','detect_correction(new_draft, existing_drafts), mark_correction_amber(draft_id)'],
    ['invoice_creator.py','create_invoice_draft, update_invoice, finalize_invoice (waybill+tax_inv+journal in 1 tx)'],
    ['normalization_service.py','normalize_amount(str→Decimal), normalize_inn(str), normalize_date(str→date)'],
],widths=[5.5,10.5])

h2('12.3 AI & Learning')
tbl(['ფაილი','ძირითადი ფუნქციები'],[
    ['ai_chat_service.py','handle_ai_chat, run_vat_calc, run_payroll_calc, run_cit_calc, run_classify_tx, run_learn_rule'],
    ['ai_service.py','call_claude_api(messages, system), call_gemini_api, build_system_prompt(tenant_context)'],
    ['llm_service.py','get_completion, get_completion_with_retry(max=3), parse_structured_response'],
    ['learning_service.py','get_patterns, update_pattern_feedback(signal), add_pattern, get_learning_stats, get_pattern_health'],
    ['pattern_decay_service.py','run_decay_service: confidence *= 0.98^age_weight, archive if <0.30'],
    ['confidence_engine.py','calculate_confidence(signals[]), combine_signals(pattern+usage+age), adjust_for_age'],
    ['context_engine.py','build_context_for_llm: recent_transactions + tenant_patterns + coa_hints'],
    ['intent_engine.py','classify_intent(message) → classify|chat|calculate|search'],
    ['feedback_service.py','record_feedback(draft_id, signal), update_pattern_from_feedback(pattern_key)'],
    ['memory_priority_engine.py','get_priority_memories(query), rank_memories_by_relevance(memories, context)'],
    ['transaction_memory_service.py','save_transaction, recall_similar(amount, counterparty), get_counterparty_history'],
],widths=[5.5,10.5])

h2('12.4 Business Logic')
tbl(['ფაილი','ძირითადი ფუნქციები'],[
    ['accounting_rules.py','50+ functions: split_vat_from_gross, build_vat_posting, build_payroll_posting, build_dividend_posting, build_depreciation_posting, gross_up_salary_from_net, estonian_cit_from_distributed'],
    ['accounting_engine.py','process_transaction(tx) → posting_batch, validate_double_entry'],
    ['payroll_service.py','calculate_payroll(gross), calculate_pit(20%), calculate_pension(2%), generate_rs_xml(employees)'],
    ['export_service.py','export_to_excel(drafts), export_to_csv(drafts), export_journal_entries(period)'],
    ['collaboration_service.py','add_comment(draft_id, text, author), assign_draft(draft_id, user), get_draft_comments'],
    ['email_invoice_service.py','parse_email_invoice(email_body), extract_invoice_from_email'],
    ['erp_import_service.py','import_from_onec(xml), map_erp_accounts(erp_code→coa), create_drafts_from_erp(data)'],
    ['erp_history_import_service.py','bulk_import_history(csv), validate_erp_data(rows), create_bulk_drafts'],
    ['tenant_service.py','create_tenant(company_name, inn, email), get_tenant(slug), update_tenant_settings'],
    ['rbac_service.py','get_user_permissions(role), check_permission(role, perm), list_users(tenant_id)'],
    ['qa_engine.py','validate_journal_entry(entry), check_accounting_rules(draft), flag_anomalies(amount, pattern)'],
    ['retry_service.py','with_retry(fn, max_attempts=3), exponential_backoff(attempt), handle_transient_errors'],
],widths=[5.5,10.5])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 13. FRONTEND PAGES
# ══════════════════════════════════════════════════════════════════════════════
h1('13. Frontend — 23 Static HTML Pages')
para('Design System: Editorial cream palette. Fonts: Inter + Instrument Serif + Geist Mono. '
     'Pattern: sidebar nav (14 items) + topbar breadcrumb + content area.')

tbl(['ფაილი','სახელი','Tabs','ძირითადი API Calls'],[
    ['approval.html','Main SPA Dashboard','Internal pages: Queue, Bank, Budget, CRM, Payroll, Reports, Export, Reconciliation + 12 more','GET /approval/queue, /dashboard/live/kpi, /bank/..., /payroll/..., /budget/..., /crm/...'],
    ['drafts.html','Drafts Hub','ყველა / მოლოდინში / დამტკ. / უარყ.','GET /approval/queue?status=, POST /approval/approve/{id}, POST /approval/reject/{id}'],
    ['documents.html','Documents Hub','Upload / OCR Queue / Email','POST /documents/upload, GET /pipeline/queue, GET /email/invoices'],
    ['ai_control.html','AI Control Panel','Patterns / Autopilot / Insights / Recommendations','GET /learning/patterns, /decision-engine/stats, /dashboard/insights, /api/ai/recommend'],
    ['settings.html','Settings Hub','General / Currency / Audit / Health / API Keys','GET /auth/me, /currency/rates, /audit-log/list, /health/'],
    ['ledger.html','Account Ledger','Single view + ?account= param','GET /reports/ledger/{code}'],
    ['trial_balance.html','Trial Balance','Single view (auto-loads)','GET /reports/trial-balance'],
    ['counterparty_ledger.html','Counterparty Ledger','Single view + INN search','GET /reports/counterparty/{inn}'],
    ['payroll_ledger.html','Payroll Ledger','Single view + year/emp filter','GET /reports/payroll'],
    ['journal.html','Journal','Card view, paginated 50/page','GET /reports/journal'],
    ['waybills.html','Waybills V2','List + stats strip + match score bar','GET /documents/waybills'],
    ['tax_invoices.html','Tax Invoices V2','List + VAT stats + rs.ge banner','GET /documents/tax-invoices'],
    ['outgoing_form.html','Outgoing Invoice Form','Goods / Service dual-mode, 2s auto-save','POST /outgoing/drafts, PATCH /outgoing/drafts/{id}, POST /outgoing/{id}/finalize'],
    ['draft_detail.html','Draft Detail','შემომავ. / გამავ. sections, journal viz','GET /approval/queue/{id}'],
    ['signup.html','Registration','Single form','POST /auth/signup'],
    ['approval_dashboard.html','Approval Dashboard (legacy)','—','—'],
    ['patterns_dashboard.html','Patterns Dashboard','—','GET /learning/patterns'],
    ['audit_dashboard.html','Audit Dashboard','—','GET /audit-log/list'],
    ['main_dashboard.html','Main Dashboard (legacy)','—','—'],
    ['settings_dashboard.html','Settings Dashboard (legacy)','—','—'],
],widths=[4.5,3,3,5.5])

h2('13.1 Design Tokens')
tbl(['CSS Variable','Hex','გამოყენება'],[
    ['--paper','#f3ecdc','Page background (cream)'],
    ['--paper-soft','#ebe2ce','Row hover, alternate backgrounds'],
    ['--paper-sunk','#e3d8bf','Chips, code blocks'],
    ['--card','#faf3e3','Panel/card surface'],
    ['--line','#dfd3b8','All borders'],
    ['--ink','#1a1512','Primary text + active sidebar bg'],
    ['--accent','#8c3c2d','CTA buttons, h2 headings'],
    ['--green','#27500A','Debit (Dr) values, positive balances'],
    ['--red','#791F1F','Credit (Cr) values, negative balances'],
    ['--amber','#7a4f00','Corrections, warnings'],
    ['--amber-bg','#FAEEDA','Correction row backgrounds'],
],widths=[4,3,9])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 14. FULL API CATALOG
# ══════════════════════════════════════════════════════════════════════════════
h1('14. API Endpoints — სრული კატალოგი')

h2('14.1 Authentication & RBAC')
tbl(['M','Path','Rate','RBAC','Description'],[
    ['POST','/auth/login','5/min','Public','Email+pw → access+refresh JWT'],
    ['POST','/auth/register','5/min','Public','Register user in existing tenant'],
    ['POST','/auth/signup','3/min','Public','Full SaaS signup → new tenant + JWT'],
    ['GET','/auth/me','—','JWT','Decode token → user info (no DB)'],
    ['POST','/auth/refresh','—','Refresh','New access token'],
    ['GET','/rbac/me','—','X-Api-Key','Current user via API key'],
    ['GET','/rbac/users','—','tenants:manage','List users'],
    ['POST','/rbac/users/create','—','tenants:manage','Create user'],
    ['GET','/tenants','—','tenants:manage','List tenants'],
    ['POST','/tenants/create','—','admin','Create tenant'],
],widths=[1.5,4.5,2,3,5])

h2('14.2 Approval')
tbl(['M','Path','Rate','RBAC','Description'],[
    ['GET','/approval/queue','—','approval:read','Status filter + pagination + pending_count'],
    ['POST','/approval/approve/{id}','30/min','approval:write','Approve (FOR UPDATE NOWAIT → 409 if locked)'],
    ['POST','/approval/reject/{id}','30/min','approval:write','Reject with reason'],
    ['POST','/approval/correct/{id}','30/min','approval:write','Correct account/entry'],
    ['GET','/approval/audit','—','audit:view','Approval history log'],
    ['POST','/approval/autopilot','—','approval:write','Manual autopilot trigger'],
    ['POST','/approval/preview','—','approval:read','Preview (no state change)'],
],widths=[1.5,4.5,2,3,5])

h2('14.3 Documents & Triangle Reconciliation')
tbl(['M','Path','Rate','Description'],[
    ['POST','/documents/upload','10/min','Upload PDF/CSV → full OCR pipeline → journal_draft'],
    ['POST','/documents/upload-waybill','20/min','Waybill OCR → waybills INSERT'],
    ['POST','/documents/upload-tax-invoice','20/min','Tax invoice OCR → tax_invoices INSERT'],
    ['POST','/documents/upload-commercial-invoice','20/min','Commercial invoice OCR'],
    ['GET','/documents/triangle-matches','—','Triangle match results with scores'],
    ['GET','/documents/waybills','—','Waybill list + match_score + status'],
    ['GET','/documents/tax-invoices','—','Tax invoice list + VAT breakdown'],
],widths=[1.5,5.5,2,7])

h2('14.4 Outgoing Invoices')
tbl(['M','Path','Rate','Description'],[
    ['POST','/outgoing/drafts','30/min','Create outgoing invoice draft'],
    ['PATCH','/outgoing/drafts/{id}','60/min','Auto-save (2s debounce client-side)'],
    ['POST','/outgoing/{id}/finalize','20/min','Finalize: waybill + tax_invoice + journal (1 TX)'],
    ['GET','/outgoing/list','30/min','List with status filter'],
    ['GET','/outgoing/{id}','—','Single invoice detail'],
    ['GET','/outgoing/{id}/pdf','10/min','PDF download'],
],widths=[1.5,5,2,7.5])

h2('14.5 Financial Reports')
tbl(['M','Path','Rate','Params'],[
    ['GET','/reports/monthly','—','—'],
    ['GET','/reports/annual','—','—'],
    ['GET','/reports/pnl','—','year, month'],
    ['GET','/reports/cashflow','—','year, month'],
    ['GET','/reports/audit-trail','—','—'],
    ['GET','/reports/ledger/{account_code}','10/min','date_from, date_to'],
    ['GET','/reports/trial-balance','10/min','date_from, date_to'],
    ['GET','/reports/counterparty/{inn}','10/min','date_from, date_to'],
    ['GET','/reports/payroll','10/min','employee_id (opt), year (opt)'],
    ['GET','/reports/journal','10/min','date (opt), limit, offset'],
],widths=[1.5,5.5,2,7])

h2('14.6 AI & Learning')
tbl(['M','Path','Rate','Description'],[
    ['POST','/api/ai/chat','—','Chat with AI (session, vector search, tenant context)'],
    ['GET','/api/ai/search','—','Vector/keyword search, top_k=5'],
    ['POST','/api/ai/vat','—','VAT calc: amount, inclusive (bool)'],
    ['POST','/api/ai/payroll','—','Payroll: gross → net, PIT 20%, pension 2%'],
    ['POST','/api/ai/dividend','—','Dividend withholding 5%'],
    ['POST','/api/ai/cit','—','Corporate income tax 15%'],
    ['POST','/api/ai/depreciation','—','Depreciation schedule'],
    ['POST','/api/ai/classify','—','text → account_code + confidence'],
    ['POST','/api/ai/learn','—','Add learned pattern'],
    ['GET','/api/ai/recommend','—','Recommendations with severity badges'],
    ['GET','/api/ai/stats','—','System stats'],
    ['GET','/learning/patterns','—','All patterns with stats'],
    ['POST','/learning/feedback','—','Submit approve/reject signal'],
    ['GET','/learning/stats','—','Learning health metrics'],
    ['GET','/decision-engine/stats','—','Queue + accuracy stats'],
    ['POST','/decision-engine/analyze','—','Force-analyze single transaction'],
],widths=[1.5,5,2,7.5])

h2('14.7 Bank, Budget, COA, Currency')
tbl(['M','Path','Description'],[
    ['POST','/bank/upload','Upload bank CSV → bank_transactions INSERT'],
    ['GET','/bank/history','CSV upload history'],
    ['GET','/posting/approved-drafts','Approved drafts ready for ERP posting'],
    ['GET','/posting/payload/{id}','Journal posting payload (1C format)'],
    ['POST','/budget/create','Create budget entry'],
    ['POST','/budget/create-annual','Bulk annual budget (multiple items)'],
    ['GET','/budget/vs-actual/{year}','Budget vs actual comparison by account'],
    ['GET','/budget/forecast/{year}','Forecast: 3-month average × 12'],
    ['GET','/budget/list/{year}','Budget entries for year'],
    ['GET','/coa/list','Chart of Accounts by category'],
    ['GET','/coa/get/{code}','Single account details'],
    ['GET','/coa/search','Search accounts by code or name'],
    ['GET','/currency/rates','Exchange rates (EUR, USD, GBP vs GEL)'],
    ['POST','/currency/convert','Convert amount between currencies'],
],widths=[1.5,5.5,9])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 15. TESTS
# ══════════════════════════════════════════════════════════════════════════════
h1('15. Test Suite')

h2('15.1 Integration Tests (tests/integration/) — 40 ფაილი')
tbl(['ფაილი','Tests','რას ამოწმებს'],[
    ['test_approval_race_condition.py','4','Concurrent approve — only one wins, second gets 409'],
    ['test_cross_tenant_isolation.py','15','Tenant A cannot read/write Tenant B data (15 scenarios)'],
    ['test_tenant_isolation.py','8','Full tenant boundary: bank, drafts, patterns, budgets'],
    ['test_tenant_middleware.py','5','X-Tenant-ID header parsing + "default" fallback'],
    ['test_document_intelligence.py','12','PDF upload → extract → classify → journal draft'],
    ['test_triangle_matcher.py','8','3-way match logic + score calculation + discrepancy detection'],
    ['test_correction_detection.py','8','Amber correction pattern detection accuracy'],
    ['test_outgoing_invoice.py','10','Draft → autosave → finalize → waybill + tax_inv + journal'],
    ['test_comment_propagation.py','6','Comment inheritance: correction draft inherits original comments'],
    ['test_autopilot_flow.py','5','Autopilot threshold enforcement (confidence, usage, age)'],
    ['test_csv_upload_flow.py','6','Bank CSV → parse → bank_transactions INSERT → dedup'],
    ['test_classification_accuracy.py','8','TBC→7720, Amazon→7810, VAT inclusive/exclusive handling'],
    ['test_duplicate_prevention.py','4','Same file_hash uploaded twice → only one draft created'],
    ['test_rbac.py','10','Permission enforcement per role (admin vs viewer vs reviewer)'],
    ['test_qa_engine.py','8','Journal entry validation: Dr=Cr balance check, account validity'],
    ['test_confidence_engine.py','6','Confidence score combination from multiple signals'],
    ['test_context_engine.py','5','LLM context building with recent transactions'],
    ['test_posting_logs.py','6','Posting log creation, retrieval, error handling'],
    ['test_workflow_layer.py','7','End-to-end: upload → approve → post to 1C'],
    ['test_invoice_to_draft.py','5','Invoice OCR → correct journal entries generated'],
    ['test_retry_service.py','4','Exponential backoff on transient DB errors'],
    ['test_memory_priority.py','4','Memory ranking by relevance score'],
    ['test_autopilot_flow.py','5','Autopilot threshold logic'],
    ['test_mock_pipeline.py','5','Pipeline tests without real LLM calls (mocked)'],
],widths=[6,2,8])

h2('15.2 Unit Tests (tests/unit/)')
tbl(['ფაილი','Tests','რას ამოწმებს'],[
    ['test_approval_queue.py','6','Queue filtering, pagination, status counts'],
    ['test_decision_engine.py','8','Pattern matching, confidence calc, fuzzy match'],
    ['test_georgia_pack.py','6','VAT 18%, PIT 20%, CIT 15%, pension 2% calculations'],
    ['test_canonical_validators.py','5','INN validation format, date normalization'],
    ['test_learning_health.py','4','Pattern decay rates, health thresholds'],
    ['test_version.py','2','GET /version endpoint returns correct version'],
    ['test_health.py','3','GET /health/ → DB connection check'],
    ['test_preview_service.py','5','Posting preview dry-run accuracy'],
],widths=[5.5,2.5,8])

h2('15.3 Test Fixtures (tests/fixtures/)')
tbl(['PDF File','Document Type'],[
    ['advance_payment.pdf','Advance payment invoice (supplier prepayment)'],
    ['invoice_buyer_it_services.pdf','IT services invoice — buyer perspective'],
    ['invoice_foreign.pdf','Foreign supplier invoice (non-Georgian)'],
    ['invoice_seller_education.pdf','Education services outgoing invoice'],
    ['office_rent_receipt.pdf','Office rent payment receipt'],
    ['saas_subscription.pdf','SaaS subscription invoice (foreign)'],
    ['utility_electricity.pdf','Utility bill — electricity (Energo-Pro)'],
],widths=[6,10])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 16. SECURITY IMPROVEMENTS
# ══════════════════════════════════════════════════════════════════════════════
h1('16. Security — ყველა გაუმჯობესება')
tbl(['#','პრობლემა','გადაწყვეტა','ფაილი/Migration'],[
    ['1','Race condition: SELECT + UPDATE ცალ-ცალკე → concurrent requests approve ერთ draft-ს ორჯერ','SELECT FOR UPDATE NOWAIT + HTTP 409 DRAFT_LOCKED response','approval_service.py, correct_draft_service.py, routes_approval.py'],
    ['2','Cross-tenant data leak: budget queries-ში (tenant_id IS NULL OR ...) → NULL rows ყველა tenant-ს ეჩვენება','WHERE tenant_id=%s strict filter (NULL rows invisible)','routes_budget.py'],
    ['3','Google Fonts blocked by CSP → UI fonts not loading','style-src + fonts.googleapis.com, font-src + fonts.gstatic.com','app/api/security.py'],
    ['4','learned_rules.json on filesystem → unprotected, survives container restart but not scaling','Migrate to learning_patterns PostgreSQL table (RLS protected)','knowledge_loader.py'],
    ['5','No rate limiting on critical endpoints → brute force / DoS possible','@limiter.limit() on /auth/login (5/min), /approval/* (30/min), /reports/* (10/min), /ocr/* (5/min)','Multiple routes files'],
    ['6','Single-tenant app → no data isolation guarantee','Full multi-tenant: tenant_id in JWT + 3-layer isolation','001_multi_tenant_schema.sql'],
    ['7','No PostgreSQL RLS → app code bug = data leak','Migration 002: RLS on all tenant tables via GUC','002_row_level_security.sql'],
    ['8','No XSS protection in dynamic HTML','esc() function in all JS: .replace(/&/g,"&amp;").replace(/</g,"&lt;")','All static HTML pages'],
],widths=[0.8,4.5,4.5,6.2])
PB()

# ══════════════════════════════════════════════════════════════════════════════
# 17. KNOWLEDGE BASE
# ══════════════════════════════════════════════════════════════════════════════
h1('17. Knowledge Base (app/knowledge/)')
tbl(['ფაილი','შინაარსი'],[
    ['chart_of_accounts.py','CHART_OF_ACCOUNTS dict (1000–9999 Georgian COA), ACCA_STANDARDS, get_account(code), search_accounts(query)'],
    ['tax_rules.py','TAX_RULES dict, calculate_vat(amount, inclusive), calculate_payroll(gross), calculate_cit(profit), calculate_withholding(amount, type), calculate_depreciation(cost, residual, life)'],
    ['knowledge_loader.py','_load_files() — JSONL KB files to memory; _load_learned() — DB first, JSON fallback; _load_learned_from_db() — learning_patterns table; learn_new_rule(key, code) — DB only; _save_learned() — no-op (deprecated); migrate_json_to_db() — one-time migration; get_stats() — KB statistics'],
    ['journal_builder.py','build_journal_from_text(text, amount) — text classification → journal entries; classify_transaction(text) → (account_code, confidence); search_knowledge(q) → matching KB entries; get_context_for_llm(transaction) → structured context'],
],widths=[4.5,11.5])

h1('18. Deployment')
tbl(['Step','Command'],[
    ['1. Code commit','git commit -m "feat: ..."'],
    ['2. Docker build','gcloud builds submit --tag europe-west1-docker.pkg.dev/project-1e145fd0-c30e-4aac-a34/cloud-run-source-deploy/fastapi-run:latest'],
    ['3. Deploy','gcloud run deploy fastapi-run --image ... --region europe-west1 --project project-1e145fd0-c30e-4aac-a34'],
    ['4. Verify','Service URL: https://fastapi-run-226875230147.europe-west1.run.app'],
    ['Current Revision','fastapi-run-00073-ctt (100% traffic)'],
    ['Build Time','~3 minutes (Cloud Build)'],
    ['Deploy Time','~1 minute (Cloud Run)'],
],widths=[3,13])

# ══════════════════════════════════════════════════════════════════════════════
# 19. STATISTICS
# ══════════════════════════════════════════════════════════════════════════════
PB()
h1('19. სტატისტიკური სარეზიუმე')
tbl(['კომპონენტი','Count','დეტალი'],[
    ['Route Modules','62 active','+ 7 archived in /archive/'],
    ['Service Modules','54','Including decision_engine submodule'],
    ['Middleware Layers','4 + 1 inline','tenant, auth, rbac, audit_log + security headers'],
    ['API Endpoints','120+','Across 62 route modules'],
    ['Static HTML Pages','23','Sidebar + topbar + content pattern'],
    ['DB Tables','30+','Across 4 migrations + pre-existing'],
    ['DB Migrations','4','001-004, all applied to production'],
    ['Roles','5','admin, accountant, reviewer, viewer, ai_supervisor'],
    ['Permissions','20+','Colon-format: reports:read, approval:write, etc.'],
    ['Integration Tests','40 files','tests/integration/'],
    ['Unit Tests','8 files','tests/unit/'],
    ['PDF Test Fixtures','7 files','Realistic Georgian business documents'],
    ['Background Tasks','2 loops','autopilot (60s) + decay (3600s)'],
    ['AI Providers','2','Anthropic Claude (primary) + Google Gemini (fallback)'],
    ['Python LOC (est.)','22,000+','Routes + Services + Middleware + Knowledge'],
    ['JavaScript LOC (est.)','8,000+','Inline in 23 HTML pages'],
    ['Total Commits (session)','15','All on main branch, all deployed'],
    ['New Files (session)','30+','HTML pages, services, tests, migrations'],
    ['Modified Files (session)','11','Routes, services, security, middleware'],
],widths=[5.5,3,7.5])

# footer
p_f = doc.add_paragraph()
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_f.add_run(
    f'Bridge Hub Financial OS · სრული ტექნიკური აუდიტი · '
    f'{datetime.date.today().strftime("%d.%m.%Y")} · '
    'შედგენილია Claude Sonnet 4.6-ის მიერ'
)
r_f.font.name = 'Calibri'; r_f.font.size = Pt(9); r_f.italic = True
r_f.font.color.rgb = RGBColor(0xa8,0x98,0x80)

out = r'C:\Users\Acer\fastapi-run\fastapi-run\BridgeHub_Full_Audit.docx'
doc.save(out)
print(f'Saved: {out}')
