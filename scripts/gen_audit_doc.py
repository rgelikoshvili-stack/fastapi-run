import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0x76)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2e, 0x86, 0xab)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for r in p.runs: r.bold = True

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def tbl_start(headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    return t

def tbl_row(t, cells):
    r = t.add_row()
    for i, v in enumerate(cells):
        r.cells[i].text = v

def mono(text):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)

# ── TITLE ──
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('Bridge Hub — Financial OS')
r.bold = True; r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.add_run('სრული სისტემური აუდიტი, არქიტექტურა და გასწორებული პრობლემები').font.size = Pt(14)

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.add_run('თარიღი: 2026-04-26  |  Revision: fastapi-run-00124').font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════
h1('1. სისტემის მიმოხილვა')
para('Bridge Hub — ქართული კომპანიებისთვის SaaS Financial OS. Google Cloud Run + PostgreSQL. AI, OCR, Bank Sync, Approval Workflow.')

h2('1.1 ტექნოლოგიური სტეკი')
t = tbl_start(['კომპონენტი', 'ტექნოლოგია'])
for row in [
    ('Backend', 'FastAPI Python 3.11'),
    ('Database', 'PostgreSQL 14 — GCP Cloud SQL (35.192.214.120)'),
    ('Hosting', 'Google Cloud Run — europe-west1'),
    ('Frontend', 'Vanilla JS + HTML5 (static files)'),
    ('Auth', 'JWT HS256 — 24h access / 7d refresh'),
    ('AI/LLM', 'Claude Anthropic + GPT-4o OpenAI + Gemini Google'),
    ('OCR', 'Tesseract + pdfminer + Vision LLM'),
    ('Rate Limiting', 'slowapi (SlowAPIMiddleware)'),
    ('CI/CD', 'GitHub Actions -> gcloud run deploy'),
    ('Background', 'asyncio loops: autopilot, pattern decay, email'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('1.2 URLs')
bullet('Production: https://fastapi-run-oobzrmikna-ew.a.run.app')
bullet('Static pages: /static/*.html')
bullet('Health: /health')

doc.add_page_break()

# ══════════════════════════════════════════
h1('2. არქიტექტურა')

h2('2.1 Middleware Pipeline (request execution order)')
t = tbl_start(['#', 'Middleware', 'ფუნქცია'])
for row in [
    ('1 (outer)', 'tenant_middleware', 'X-Tenant-ID header / query param -> request.state.tenant_id'),
    ('2', 'auth_middleware', 'JWT decode -> tenant_id, user_id, role (overwrites tenant_id)'),
    ('3', 'rbac_middleware', 'Role-Based Access Control — permission check per route'),
    ('4 (inner)', 'audit_log_middleware', 'All HTTP request/response logged to audit_log table'),
    ('5', 'add_security_headers', 'HSTS, CSP, X-Frame-Options, X-XSS-Protection'),
    ('6', 'SlowAPIMiddleware', 'Rate limiting via slowapi'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('2.2 Tenant Isolation')
bullet('ყველა DB query: WHERE tenant_id = %s')
bullet('resolve_tenant_id(): isdigit() -> _inn_to_tenant_id() -> lru_cache(128) -> 202192643 -> "default"')
bullet('startup migration: normalizes company_inn -> tenant_id in all tables')

h2('2.3 Document Intelligence Pipeline')
steps = [
    'POST /documents/upload',
    '  1. File size check (max 10MB)',
    '  2. SHA-256 dedup: duplicate+content -> return existing | duplicate+NULL content -> UPDATE (content_restored)',
    '  3. parse_document() -> text (pdfminer / Tesseract / Vision LLM)',
    '  4. extract_document() -> ExtractedDocument(seller, buyer, amount, date)',
    '     - LLM: Georgian keyword-aware prompt (gamyidveli/myidveli)',
    '     - Regex fallback: _extract_inn_near() with keyword anchors',
    '  5. resolve_party() -> OurRole: BUYER / SELLER / INTERNAL / FOREIGN',
    '     - _load_tenant(tenant_id): try tenant_id -> fallback company_inn',
    '  6. FOREIGN -> status=pending_human_review, Dr/Cr=????',
    '  7. BUYER/SELLER -> classify_operation() -> build_journal() -> status=pending',
    '  8. INSERT journal_drafts + INSERT processed_documents(file_content=bytes)',
]
for s in steps:
    bullet(s)

doc.add_page_break()

# ══════════════════════════════════════════
h1('3. მონაცემთა ბაზა')

h2('3.1 journal_drafts — მთავარი სამუშაო ცხრილი')
t = tbl_start(['სვეტი', 'ტიპი', 'დანიშნულება'])
for row in [
    ('id', 'SERIAL PK', 'auto increment'),
    ('tenant_id', 'TEXT DEFAULT default', 'multi-tenant isolation'),
    ('status', 'TEXT', 'pending | pending_human_review | approved | rejected | auto_approved | posted'),
    ('description', 'TEXT', 'ტრანზაქციის აღწერა'),
    ('partner', 'TEXT', 'კონტრაგენტი'),
    ('amount', 'NUMERIC', 'თანხა'),
    ('currency', 'TEXT DEFAULT GEL', 'ვალუტა'),
    ('debit_account', 'TEXT', 'Dr ანგარიში (COA კოდი) — ???? if unknown'),
    ('credit_account', 'TEXT', 'Cr ანგარიში — ???? if unknown'),
    ('confidence', 'NUMERIC', 'AI confidence 0.0 - 1.0'),
    ('our_role', 'TEXT', 'buyer | seller | internal | foreign'),
    ('source_document_id', 'INTEGER FK', 'processed_documents.id (eye button)'),
    ('is_foreign_doc', 'BOOLEAN', 'True = სხვა კომპანიის დოკუმენტი'),
    ('journal_entries', 'JSONB', 'Dr/Cr ჩანაწერები JSON array'),
    ('raw_extraction', 'JSONB', 'OCR raw result'),
    ('created_at', 'TIMESTAMPTZ', ''),
    ('updated_at', 'TIMESTAMPTZ', ''),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('3.2 processed_documents — ატვირთული ფაილები')
t = tbl_start(['სვეტი', 'ტიპი', 'დანიშნულება'])
for row in [
    ('id', 'SERIAL PK', ''),
    ('tenant_id', 'TEXT', 'isolation'),
    ('file_hash', 'TEXT UNIQUE', 'SHA-256 dedup key'),
    ('file_name', 'TEXT', 'ფაილის სახელი'),
    ('mime_type', 'TEXT', 'application/pdf | image/...'),
    ('file_content', 'BYTEA', 'ფაილის bytes — eye button preview (added in this session)'),
    ('extraction_method', 'TEXT', 'native_pdf | tesseract_pdf | vision_llm'),
    ('raw_text', 'TEXT', 'OCR text (max 10000 chars)'),
    ('extracted_data', 'JSONB', 'ExtractedDocument JSON'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('3.3 სხვა ცხრილები')
t = tbl_start(['ცხრილი', 'დანიშნულება'])
for row in [
    ('tenants', 'tenant_id, company_inn, company_name, is_vat_payer'),
    ('users', 'email, password_hash, tenant_id, role'),
    ('bank_transactions', 'Bank Sync transactions'),
    ('learning_patterns', 'AI learned rules + decay scores'),
    ('audit_log', 'All HTTP requests/responses'),
    ('counterparties', 'Vendors / customers'),
    ('chart_of_accounts', 'COA per tenant'),
    ('invoices', 'Outgoing invoices'),
    ('expenses', 'Expense records'),
    ('contracts', 'Contracts + milestones'),
    ('customers', 'CRM customers'),
    ('draft_comments', 'Comments on drafts'),
    ('chat_sessions', 'AI chat history'),
]:
    tbl_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════
h1('4. ფაილების სრული კატალოგი')

h2('4.1 Root')
t = tbl_start(['ფაილი', 'დახასიათება'])
for row in [
    ('main.py', 'FastAPI entry point. Routers, middleware, _run_db_migrations(), startup loops'),
    ('bridge_hub_knowledge.py', 'COA, tax rules, transaction classifier, LLM context. Shim -> app/knowledge/'),
    ('bridge_hub_vector_db.py', 'Vector DB (Pinecone/FAISS) semantic search'),
    ('requirements.txt', 'Python dependencies'),
    ('Dockerfile', 'Cloud Run container'),
    ('Procfile', 'gunicorn/uvicorn startup'),
    ('gen_audit_doc.py', 'This document generator'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('4.2 app/api/ — API Routes')
t = tbl_start(['ფაილი', 'Prefix', 'ფუნქცია'])
for row in [
    ('routes_approval.py', '/approval', 'queue, approve, reject, correct, batch-action, preview, stats, autopilot'),
    ('routes_documents.py', '/documents', 'upload pipeline, /file endpoint, triangle matching'),
    ('routes_auth.py', '/auth', 'login, register, refresh, password reset'),
    ('routes_posting.py', '/posting', 'apply/{id}, preview/{id}, history'),
    ('routes_ai_chat.py', '/api/ai/chat', 'Claude/GPT accounting assistant'),
    ('routes_claude_chat.py', '/api/claude', 'Claude direct API'),
    ('routes_bank_sync.py', '/bank-sync', 'Bank sync: connect, fetch, process'),
    ('routes_bank_csv.py', '/bank/csv', 'CSV upload + parse'),
    ('routes_bank_accounts.py', '/bank-accounts', 'CRUD bank accounts'),
    ('routes_dashboard.py', '/dashboard', 'KPIs: revenue, expenses, cash flow'),
    ('routes_dashboard_live.py', '/dashboard/live', 'Real-time dashboard'),
    ('routes_reports.py', '/reports', 'P&L, Balance Sheet, Cash Flow'),
    ('routes_coa.py', '/coa', 'Chart of Accounts CRUD'),
    ('routes_invoices.py', '/invoices', 'Outgoing invoices'),
    ('routes_expenses.py', '/expenses', 'Expense management'),
    ('routes_crm.py', '/crm', 'Customers, interactions, pipeline'),
    ('routes_contracts.py', '/contracts', 'Contracts + milestones'),
    ('routes_budget.py', '/budget', 'Budget planning + variance'),
    ('routes_payroll.py', '/payroll', 'Georgian payroll PIT/PAYG'),
    ('routes_tax.py', '/tax', 'VAT, CIT, withholding tax'),
    ('routes_reconciliation.py', '/reconciliation', 'Bank vs ledger reconciliation'),
    ('routes_export.py', '/export', 'Excel/CSV export'),
    ('routes_learning.py', '/learning', 'AI learning patterns'),
    ('routes_notifications.py', '/notifications', 'Push notifications'),
    ('routes_audit.py', '/audit', 'Audit trail'),
    ('routes_audit_log.py', '/audit-log', 'Raw audit log viewer'),
    ('routes_tenants.py', '/tenants', 'Tenant settings'),
    ('routes_rbac.py', '/rbac', 'Role management'),
    ('routes_1c.py', '/1c', '1C:Enterprise export'),
    ('routes_balance_ge.py', '/balance-ge', 'Balance.ge integration'),
    ('routes_transaction_ai.py', '/api/ai/transaction', 'AI transaction classification'),
    ('routes_health.py', '/health', 'Health check'),
    ('routes_security.py', '/security', 'Security audit'),
    ('routes_search.py', '/search', 'Full-text search'),
    ('routes_email_invoice.py', '/email-invoice', 'Email invoice intake'),
    ('routes_email_collector.py', '/email-collector', 'IMAP email polling'),
    ('routes_collaboration.py', '/collaboration', 'Draft comments multi-user'),
    ('routes_patterns.py', '/patterns', 'Classification pattern explorer'),
    ('routes_decision_engine.py', '/decision', 'Rule-based decisions'),
    ('routes_currency.py', '/currency', 'NBG exchange rates'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('4.3 app/api/middleware/')
t = tbl_start(['ფაილი', 'ფუნქცია'])
for row in [
    ('auth_middleware.py', 'JWT decode -> request.state.tenant_id, user_id, role'),
    ('tenant_middleware.py', 'X-Tenant-ID header fallback (overwritten by auth_middleware)'),
    ('rbac_middleware.py', 'Route permission enforcement'),
    ('audit_log_middleware.py', 'All requests -> audit_log table'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('4.4 app/api/services/ — Business Logic')
t = tbl_start(['სერვისი', 'დახასიათება'])
for row in [
    ('approval_service.py', 'approve/reject/autopilot. SELECT FOR UPDATE — race condition protection'),
    ('party_resolver.py', 'BUYER/SELLER/INTERNAL/FOREIGN detection. company_inn fallback fix (this session)'),
    ('document_extractor.py', 'LLM + regex extraction. Keyword anchors: gamyidveli/myidveli (fixed this session)'),
    ('document_parser.py', 'PDF/image -> text: native_pdf, tesseract_pdf, vision_llm'),
    ('posting_preview_service.py', 'Dr/Cr impact preview without DB write'),
    ('posting_service.py', 'Draft -> journal_entries posting'),
    ('ai_chat_service.py', 'Chat logic, context injection, draft creation'),
    ('learning_service.py', 'Pattern learning + run_decay_service()'),
    ('pattern_decay_service.py', 'Pattern confidence decay over time'),
    ('confidence_engine.py', 'Multi-factor confidence scoring'),
    ('auth_service.py', 'JWT create/verify, login, refresh'),
    ('journal_service.py', 'Journal validation, Georgian INN validate'),
    ('ledger_service.py', 'Account ledger queries'),
    ('correct_draft_service.py', 'Draft correction workflow'),
    ('triangle_matcher.py', 'Waybill + tax invoice + payment 3-way match'),
    ('doc_journal_builder.py', 'ExtractedDocument -> Journal entries'),
    ('operation_classifier.py', 'Document -> operation category'),
    ('email_collector.py', 'IMAP polling, invoice extraction'),
    ('export_service.py', 'Excel/CSV generation'),
    ('payroll_service.py', 'Georgian payroll: PIT 20%, PAYG 2%'),
    ('llm_service.py', 'LLM abstraction: OpenAI + Anthropic + Gemini'),
    ('tenant_service.py', 'Tenant CRUD'),
    ('insights_service.py', 'AI financial insights'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('4.5 app/api/tenant_context.py')
para('resolve_tenant_id(value) — ნორმალიზაცია:')
bullet('value.isdigit() -> _inn_to_tenant_id(inn) -> DB: SELECT tenant_id FROM tenants WHERE company_inn=%s')
bullet('lru_cache(128) — პირველი DB hit-ის შემდეგ cached')
bullet('202192643 -> "default" (Alte University)')

doc.add_paragraph()
h2('4.6 app/knowledge/')
t = tbl_start(['ფაილი', 'შინაარსი'])
for row in [
    ('chart_of_accounts.py', 'Georgian COA: 1xxx-9xxx accounts'),
    ('tax_rules.py', 'VAT 18%, PIT 20%, PAYG 2%, CIT 15%, withholding'),
    ('journal_builder.py', 'build_journal_from_text(), classify_transaction()'),
    ('knowledge_loader.py', '_load_learned(), learn_new_rule(), get_stats()'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('4.7 app/policy/')
t = tbl_start(['ფაილი', 'შინაარსი'])
for row in [
    ('localization/georgia_pack.py', 'calculate_vat, calculate_payroll, calculate_cit, calculate_withholding'),
    ('tax_constants.py', 'TAX_RATES dict'),
    ('audit_rules.py', 'Audit check rules'),
    ('control_rules.py', 'Internal control rules'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('4.8 static/ — Frontend Pages')
t = tbl_start(['HTML', 'გვერდი'])
for row in [
    ('approval.html', 'Approval Queue — drafts, approve/reject, eye button, posting preview, footer actions'),
    ('documents.html', 'Documents Hub — upload, OCR scan, email intake, triangle upload'),
    ('drafts.html', 'Draft list with filters'),
    ('main_dashboard.html', 'Dashboard KPIs + charts'),
    ('journal.html', 'Journal entries'),
    ('ledger.html', 'Account Ledger'),
    ('trial_balance.html', 'Trial Balance'),
    ('waybills.html', 'Waybills list'),
    ('tax_invoices.html', 'Tax invoices'),
    ('counterparty_ledger.html', 'Counterparty Ledger'),
    ('settings.html', 'System settings'),
    ('audit_dashboard.html', 'Audit log viewer'),
    ('patterns_dashboard.html', 'AI patterns explorer'),
    ('payroll_ledger.html', 'Payroll management'),
    ('outgoing_form.html', 'Outgoing invoice creator'),
    ('signup.html', 'Registration'),
    ('approval_dashboard.html', 'Approval statistics'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('4.9 app/engines/')
t = tbl_start(['Engine', 'ფუნქცია'])
for row in [
    ('accounting_engine.py', 'Double-entry validation'),
    ('audit_engine.py', 'Automated audit checks'),
    ('finance_engine.py', 'Financial KPI calculations'),
    ('gaas_engine.py', 'GaaS compliance (Georgian Accounting Standards)'),
    ('reconciliation_engine.py', 'Bank vs ledger reconciliation'),
    ('strategy_engine.py', 'Financial strategy analysis'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('4.10 scripts/ — Migration Scripts')
t = tbl_start(['Script', 'დანიშნულება'])
for row in [
    ('run_001_migration.py', 'Initial schema migration'),
    ('run_002_migration.py', 'Second schema migration'),
    ('run_tenant_migration.py', 'Tenant table setup'),
    ('run_users_migration.py', 'Users table + roles'),
    ('run_audit_migration.py', 'audit_log table'),
    ('find_missing_tenant_filters.py', 'Audit: find queries without tenant_id filter'),
    ('run_bank_files_migration.py', 'Bank files table migration'),
    ('run_contracts_migration.py', 'Contracts table migration'),
    ('run_indexer.py', 'Vector DB indexing'),
]:
    tbl_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════
h1('5. გამოსწორებული პრობლემები — 2026-04-26 სესია')

bugs = [
    ('BUG-01', 'Duplicate X button — correction modal',
     'openEditModal() ამატებდა closeModal ✕ ბუტონს modal-title innerHTML-ში. panel-close ✕ უკვე არსებობდა.',
     'modal-title.innerHTML = plain title text only. ✕ btn modal header-ში ამოღებულია.',
     'static/approval.html'),

    ('BUG-02', 'Three-dot (triple-dot) menu in preview modal',
     'Draft preview-ის modal-title-ში iyo kreba dropdown (three-dot). UI-ში X-ს გვerdze gaCanda.',
     'Three-dot dropdown სრულად ამოღებულია. Footer-ში პირდაპირ: Redaktireba / Uarqopa / Damtkiceba.',
     'static/approval.html'),

    ('BUG-03', 'Race condition: posting preview overwrites edit modal',
     'showPostingPreview() async fetch dasruldeboda openEditModal()-ის shemdes da modal content-s gadawerdad.',
     '_modalToken counter: ++token before fetch, check after. If token changed -> ignore response.',
     'static/approval.html'),

    ('BUG-04', 'file_content column missing in processed_documents',
     'Documents Hub upload iZlevoda DB error: column file_content does not exist.',
     'ALTER TABLE processed_documents ADD COLUMN IF NOT EXISTS file_content BYTEA — on production DB.',
     'main.py + production DB'),

    ('BUG-05', 'Eye button: document not found (tenant_id mismatch)',
     'showDraftDocument() -> /posting/preview/{id}: source_document_id=null. Draft tenant_id=202192643, JWT=default.',
     '(a) DB records updated: 202192643 -> default. (b) resolve_tenant_id() INN normalization.',
     'app/api/tenant_context.py'),

    ('BUG-06', 'Dedup blocks re-upload when file_content is NULL',
     'Old upload with file_content=NULL blocked re-upload via hash dedup. User could not restore bytes.',
     'Dedup query: SELECT file_content IS NOT NULL. If NULL -> UPDATE file_content (status: content_restored).',
     'app/api/routes_documents.py'),

    ('BUG-07', 'Foreign documents auto-rejected (rejected_foreign)',
     'FOREIGN role docs got status=rejected_foreign -> filtered out of Approval Queue. User could not review.',
     'FOREIGN -> status=pending_human_review. description/partner/reason populated. source_document_id linked.',
     'app/api/routes_documents.py'),

    ('BUG-08', 'party_resolver: FOREIGN misclassification (tenant INN lookup fail)',
     '_load_tenant(tenant_id=202192643) -> not found (DB has tenant_id=default). Empty INN -> FOREIGN.',
     '_load_tenant() fallback: WHERE company_inn=%s. 202192643 finds Alte University -> correct BUYER/SELLER.',
     'app/api/services/party_resolver.py'),

    ('BUG-09', 'tenant_id inconsistency: company_inn vs tenant slug',
     'JWT tenant_id=202192643 (INN) but DB tenant_id=default. All queries by one tenant failed for the other.',
     'resolve_tenant_id(): isdigit() -> _inn_to_tenant_id(lru_cache) -> normalizes. Startup migration syncs rows.',
     'app/api/tenant_context.py + main.py'),

    ('BUG-10', 'Document extractor: seller/buyer swap',
     '_regex_extract() assigned first INN in text as seller. PCShop invoice: buyer INN (ours) appeared first.',
     '_extract_inn_near() with Georgian keyword anchors. LLM prompt improved with gamyidveli/myidveli guidance.',
     'app/api/services/document_extractor.py'),

    ('BUG-11', 'GitHub Actions Checkout failure',
     'actions/checkout@v4 and @v3 both fail at Checkout step. Root cause: possibly repo/GCP permissions.',
     'Workaround: direct gcloud run deploy from local machine. CI/CD still broken — needs investigation.',
     '.github/workflows/deploy.yml'),
]

for bug_id, title, problem, fix, files in bugs:
    p = doc.add_paragraph()
    r = p.add_run(f'{bug_id}  {title}')
    r.bold = True
    r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    doc.add_paragraph(f'პრობლემა:  {problem}')
    doc.add_paragraph(f'გამოსწორება:  {fix}')
    fp = doc.add_paragraph(f'ფაილი: {files}')
    if fp.runs:
        fp.runs[0].font.name = 'Courier New'
        fp.runs[0].font.size = Pt(9)
        fp.runs[0].font.color.rgb = RGBColor(0x14, 0x6b, 0x3a)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════
h1('6. API Endpoints — სრული სია')

t = tbl_start(['Method', 'Endpoint', 'Auth', 'ფუნქცია'])
for row in [
    ('POST', '/auth/login', 'Public', 'JWT token'),
    ('POST', '/auth/register', 'Public', 'User registration'),
    ('GET', '/health', 'Public', 'Health check'),
    ('POST', '/documents/upload', 'JWT', 'Document upload + OCR'),
    ('GET', '/documents/{id}/file', 'JWT', 'File bytes (eye button)'),
    ('GET', '/documents/{id}', 'JWT', 'Document metadata'),
    ('GET', '/approval/queue', 'JWT', 'Pending drafts list'),
    ('POST', '/approval/approve/{id}', 'JWT+limit', 'Approve draft'),
    ('POST', '/approval/reject/{id}', 'JWT+limit', 'Reject draft'),
    ('POST', '/approval/correct/{id}', 'JWT+limit', 'Correct with accounts'),
    ('PATCH', '/approval/draft/{id}', 'JWT', 'Update draft fields'),
    ('POST', '/approval/batch-action', 'JWT+limit', 'Bulk approve/reject'),
    ('POST', '/approval/preview', 'JWT', 'Draft preview'),
    ('GET', '/approval/stats', 'JWT', 'Queue statistics'),
    ('POST', '/approval/autopilot', 'JWT', 'Auto-approve high-confidence'),
    ('GET', '/approval/audit', 'JWT', 'Approval audit log'),
    ('GET', '/posting/preview/{id}', 'JWT', 'Posting impact preview'),
    ('POST', '/posting/apply/{id}', 'JWT', 'Post draft to journal'),
    ('POST', '/api/ai/chat', 'JWT', 'AI chat assistant'),
    ('GET', '/dashboard/live', 'JWT', 'Real-time KPIs'),
    ('GET', '/reports/pl', 'JWT', 'Profit & Loss'),
    ('GET', '/reports/balance-sheet', 'JWT', 'Balance Sheet'),
    ('GET', '/reports/cash-flow', 'JWT', 'Cash Flow'),
    ('GET', '/coa', 'JWT', 'Chart of Accounts'),
    ('POST', '/coa', 'JWT', 'Add COA entry'),
    ('GET', '/bank-sync/transactions', 'JWT', 'Bank transactions'),
    ('POST', '/bank-sync/fetch', 'JWT', 'Fetch bank data'),
    ('POST', '/bank/csv/upload', 'JWT', 'Upload CSV statement'),
    ('GET', '/tenants/me', 'JWT', 'Current tenant info'),
    ('PATCH', '/tenants/me', 'JWT', 'Update tenant settings'),
]:
    tbl_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════
h1('7. უსაფრთხოება')

h2('7.1 Authentication')
bullet('JWT HS256 signed with JWT_SECRET env variable')
bullet('Access Token TTL: 24 hours')
bullet('Refresh Token TTL: 7 days')
bullet('Roles: admin, accountant, viewer, auditor, financial_lead')

h2('7.2 Rate Limiting (slowapi)')
t = tbl_start(['Endpoint', 'Limit'])
for row in [
    ('/auth/login', '5/minute — brute force protection'),
    ('/approval/approve, /reject, /correct', '30/minute per endpoint'),
    ('/documents/upload', '10/minute'),
    ('/approval/batch-action', '30/minute'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('7.3 XSS Protection')
bullet('esc() function — all user input HTML-escaped before DOM insertion')
bullet('Content-Security-Policy header')
bullet('X-Frame-Options: DENY')
bullet('X-XSS-Protection: 1; mode=block')
bullet('Strict-Transport-Security (HSTS)')

h2('7.4 Tenant Isolation')
bullet('All DB queries: WHERE tenant_id = %s')
bullet('JWT carries tenant_id — validated server-side on every request')
bullet('company_inn -> tenant_id normalization at resolve_tenant_id()')
bullet('Startup migration: normalizes any INN-keyed records to tenant slug')

doc.add_page_break()

# ══════════════════════════════════════════
h1('8. Deployment')

h2('8.1 Cloud Run Parameters')
t = tbl_start(['Parameter', 'Value'])
for row in [
    ('Service', 'fastapi-run'),
    ('Region', 'europe-west1'),
    ('Latest revision', 'fastapi-run-00124-mrv'),
    ('URL', 'https://fastapi-run-oobzrmikna-ew.a.run.app'),
    ('Platform', 'managed (serverless)'),
    ('Unauthenticated', 'allowed (JWT app-level)'),
]:
    tbl_row(t, row)

doc.add_paragraph()
h2('8.2 Deploy Command (manual)')
mono('gcloud run deploy fastapi-run --source . --region europe-west1 --platform managed --allow-unauthenticated --quiet')

h2('8.3 Required Secrets')
bullet('DATABASE_URL — postgresql://postgres:...@35.192.214.120/bridgehub')
bullet('JWT_SECRET — signing key')
bullet('ANTHROPIC_API_KEY — Claude')
bullet('OPENAI_API_KEY — GPT-4o')
bullet('GOOGLE_API_KEY — Gemini')
bullet('GCP_PROJECT_ID — GCP project')
bullet('GCP_SA_KEY — Service Account JSON (for CI/CD)')

doc.add_page_break()

# ══════════════════════════════════════════
h1('9. Tests')

h2('Unit Tests — tests/unit/')
for f in [
    'test_approval_queue.py — approval workflow',
    'test_version.py — app version',
    'test_health.py — /health endpoint',
    'test_learning_health.py — learning patterns',
    'test_georgia_pack.py — VAT/payroll calc',
    'test_decision_engine.py — rule engine',
    'test_canonical_validators.py — data validation',
    'test_preview_service.py — posting preview',
]:
    bullet(f)

h2('Integration Tests — tests/integration/')
for f in [
    'test_tenant_isolation.py',
    'test_cross_tenant_isolation.py',
    'test_approval_race_condition.py',
    'test_document_intelligence.py',
    'test_invoice_to_draft.py',
    'test_triangle_matcher.py',
    'test_classification_accuracy.py',
    'test_duplicate_prevention.py',
    'test_chat_intent.py + test_chat_execute_flow.py',
]:
    bullet(f)

h2('CI/CD Status')
bullet('GitHub Actions: .github/workflows/deploy.yml')
bullet('test job: pytest tests/unit/ -v')
bullet('deploy job: gcloud run deploy')
bullet('smoke test: curl /health -> HTTP 200')
bullet('STATUS: BROKEN — Checkout step fails (permissions issue)')
bullet('Workaround: manual gcloud run deploy from local')

doc.add_page_break()

# ══════════════════════════════════════════
h1('10. ცნობილი პრობლემები და შემდეგი ნაბიჯები')

h2('10.1 Open Issues')
bullet('GitHub Actions broken — root cause: GCP_SA_KEY scope or repo permissions')
bullet('Tesseract OCR accuracy for Georgian text — needs fine-tuning')
bullet('RBAC: 3 separate implementations (authz.py, rbac.py, compliance/rbac.py) — needs consolidation')
bullet('learned_rules.json legacy file — JSON fallback still in code, needs DB-only migration')
bullet('lru_cache invalidation — _inn_to_tenant_id cache not cleared on tenant update')
bullet('file_content in BYTEA — large files may hit Cloud Run memory limit; consider GCS')

h2('10.2 Next Steps')
bullet('Fix GitHub Actions — investigate GCP_SA_KEY permissions')
bullet('Triangle matching — full 3-way auto-match waybill + tax invoice + payment')
bullet('Email auto-intake — IMAP -> auto-upload + draft creation')
bullet('Multi-currency — USD/EUR/GBP full support with NBG rates')
bullet('RBAC consolidation — single source of truth')
bullet('learned_rules.json -> DB only')

# ── FOOTER ──
doc.add_page_break()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run('Bridge Hub Financial OS  |  Confidential Internal Document  |  2026-04-26').font.size = Pt(9)

doc.save(r'C:/Users/Acer/Desktop/BridgeHub_Audit_2026.docx')
print('DONE: C:/Users/Acer/Desktop/BridgeHub_Audit_2026.docx')
