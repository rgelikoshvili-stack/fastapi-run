"""Generate Bridge Hub Statistics Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

DARK = RGBColor(0x1a, 0x1a, 0x2e)
ACCENT = RGBColor(0x8b, 0x5e, 0x3c)
AI_COLOR = RGBColor(0x2d, 0x6a, 0x4f)
HUB_COLOR = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x6b, 0x70, 0x80)
WHITE = RGBColor(0xff, 0xff, 0xff)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def bar(pct, width=18):
    filled = round(pct / 100 * width)
    return "\u2588" * filled + "\u2591" * (width - filled)


# ── TITLE ─────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BRIDGE HUB")
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Financial OS  —  Bridge Hub vs AI Statistics")
r2.font.size = Pt(14)
r2.font.color.rgb = ACCENT

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("2026  |  r.gelikoshvili@gmail.com")
r3.font.size = Pt(10)
r3.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph("\u2500" * 55)
doc.add_paragraph()


# ── HELPER: table ─────────────────────────────────────────────
def make_table(data, col_widths, header_bg="1a1a2e", alt_row=False):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = "Table Grid"
    for i, row_data in enumerate(data):
        row = t.rows[i]
        for j, txt in enumerate(row_data):
            c = row.cells[j]
            c.width = col_widths[j]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp = c.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = pp.add_run(str(txt))
            rr.font.size = Pt(9)
            if i == 0:
                rr.font.bold = True
                rr.font.color.rgb = WHITE
                set_cell_bg(c, header_bg)
            else:
                if alt_row and i % 2 == 1:
                    set_cell_bg(c, "f9f9f9")
    return t


def add_h1(text, color=DARK):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.color.rgb = color
        r.font.bold = True


# ── SECTION 1 — SUMMARY ───────────────────────────────────────
add_h1("1.  Overall Statistics")

summary = [
    ["Metric", "Bridge Hub Code", "AI Functions", "Total"],
    ["Route files", "54", "13", "67"],
    ["Service files", "32", "20", "52"],
    ["Functions (routes)", "277", "61", "338"],
    ["Functions (services)", "197", "107", "304"],
    ["Total functions", "474", "168", "642"],
    ["Share", "73.8 %", "26.2 %", "100 %"],
]

t1 = doc.add_table(rows=len(summary), cols=4)
t1.style = "Table Grid"
widths = [Cm(5.5), Cm(4), Cm(4), Cm(3)]

for i, row_data in enumerate(summary):
    row = t1.rows[i]
    for j, txt in enumerate(row_data):
        c = row.cells[j]
        c.width = widths[j]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(txt)
        rr.font.size = Pt(9.5)
        if i == 0:
            rr.font.bold = True
            rr.font.color.rgb = WHITE
            set_cell_bg(c, "1a1a2e")
        elif i == len(summary) - 1:
            rr.font.bold = True
            if j == 1:
                rr.font.color.rgb = HUB_COLOR
                set_cell_bg(c, "dde8ff")
            elif j == 2:
                rr.font.color.rgb = AI_COLOR
                set_cell_bg(c, "d4edda")
        else:
            if j == 1:
                rr.font.color.rgb = HUB_COLOR
                set_cell_bg(c, "eef2ff")
            elif j == 2:
                rr.font.color.rgb = AI_COLOR
                set_cell_bg(c, "e8f5ee")

doc.add_paragraph()


# ── SECTION 2 — BAR CHARTS ────────────────────────────────────
add_h1("2.  Percentage Distribution (Visual)")

bars = [
    ["Category", "Bridge Hub Code", "AI Functions"],
    ["Total functions (642)",
     bar(73.8) + "  73.8 %",
     bar(26.2) + "  26.2 %"],
    ["Route files (67)",
     bar(80.6) + "  80.6 %",
     bar(19.4) + "  19.4 %"],
    ["Service files (52)",
     bar(61.5) + "  61.5 %",
     bar(38.5) + "  38.5 %"],
]

t2 = doc.add_table(rows=len(bars), cols=3)
t2.style = "Table Grid"
w2 = [Cm(4.5), Cm(6.5), Cm(5.5)]

for i, row_data in enumerate(bars):
    row = t2.rows[i]
    for j, txt in enumerate(row_data):
        c = row.cells[j]
        c.width = w2[j]
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
        rr = pp.add_run(txt)
        rr.font.size = Pt(9) if i > 0 else Pt(9)
        if i == 0:
            rr.font.bold = True
            rr.font.color.rgb = WHITE
            set_cell_bg(c, "1a1a2e")
        elif j == 1:
            rr.font.color.rgb = HUB_COLOR
        elif j == 2:
            rr.font.color.rgb = AI_COLOR

doc.add_paragraph()


# ── SECTION 3 — BRIDGE HUB DOMAINS ───────────────────────────
add_h1("3.  Bridge Hub — Core Code Domains")

hub_domains = [
    ["Domain", "Functions", "Share", "Description"],
    ["Auth / RBAC / Security", "34", "17.3 %", "JWT, roles, tenant isolation, API keys"],
    ["Notifications / Webhooks", "31", "15.7 %", "Real-time push, WebSocket, collaboration"],
    ["CRM / Invoice / Contracts", "30", "15.2 %", "Customers, invoices, contract tracking"],
    ["Reports / Export / Balance", "29", "14.7 %", "P&L, Trial Balance, PDF/Excel/CSV export"],
    ["Approval / Posting", "26", "13.2 %", "Draft queue, autopilot, batch approve"],
    ["Payroll / Tax", "16", "8.1 %", "PIT 20%, PAYG 2%, CIT 15%, salary calc"],
    ["Bank / Sync", "13", "6.6 %", "TBC/BOG sync, CSV import, bank accounts"],
    ["1C / ERP Integration", "8", "4.1 %", "1C export, ERP history, connectors"],
    ["Dashboard", "7", "3.6 %", "KPI cards, financial metrics, live data"],
    ["Budget", "5", "2.5 %", "Annual planning, budget vs actual"],
]

t3 = doc.add_table(rows=len(hub_domains), cols=4)
t3.style = "Table Grid"
w3 = [Cm(4.5), Cm(2.2), Cm(2.2), Cm(7.5)]

for i, row_data in enumerate(hub_domains):
    row = t3.rows[i]
    for j, txt in enumerate(row_data):
        c = row.cells[j]
        c.width = w3[j]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
        rr = pp.add_run(txt)
        rr.font.size = Pt(8.5)
        if i == 0:
            rr.font.bold = True
            rr.font.color.rgb = WHITE
            set_cell_bg(c, "1a1a2e")
        else:
            if j == 0:
                rr.font.bold = True
                rr.font.color.rgb = HUB_COLOR
            elif j == 2:
                rr.font.bold = True
            elif j == 3:
                rr.font.color.rgb = GRAY
            if i % 2 == 0:
                set_cell_bg(c, "eef2ff")

doc.add_paragraph()


# ── SECTION 4 — AI CATEGORIES ────────────────────────────────
add_h1("4.  AI Functions — By Category", AI_COLOR)

ai_cats = [
    ["AI Component", "Functions", "Share", "What it does"],
    ["Chat / Conversation (Claude)", "20", "19.0 %", "Natural language Q&A, history, role profiles"],
    ["LLM Service (Claude/GPT/Gemini)", "13", "12.4 %", "API calls, cost logging, conversation history"],
    ["Email AI (parse + classify)", "13", "12.4 %", "Email invoice parsing, auto-draft creation"],
    ["Learning / Patterns", "10", "9.5 %", "Pattern learning, decay, memory priority"],
    ["AI Processor / Service", "8", "7.6 %", "Context builder, pipeline, AI service layer"],
    ["AI Classification", "8", "7.6 %", "Transaction classifier, decision engine, intent"],
    ["OCR / Documents", "5", "4.8 %", "PDF/image OCR, LLM extraction, INN validation"],
    ["Transaction Memory", "5", "4.8 %", "Memory search, partner resolution, scoring"],
    ["ERP Memory", "3", "2.9 %", "ERP history learning, pattern memory"],
    ["Confidence Engine", "1", "1.0 %", "Confidence scoring for classifications"],
    ["QA Engine (Gemini)", "1", "1.0 %", "Q&A over accounting knowledge base"],
]

t4 = doc.add_table(rows=len(ai_cats), cols=4)
t4.style = "Table Grid"
w4 = [Cm(5), Cm(2.2), Cm(2.2), Cm(7)]

for i, row_data in enumerate(ai_cats):
    row = t4.rows[i]
    for j, txt in enumerate(row_data):
        c = row.cells[j]
        c.width = w4[j]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
        rr = pp.add_run(txt)
        rr.font.size = Pt(8.5)
        if i == 0:
            rr.font.bold = True
            rr.font.color.rgb = WHITE
            set_cell_bg(c, "2d6a4f")
        else:
            if j == 0:
                rr.font.bold = True
                rr.font.color.rgb = AI_COLOR
            elif j == 2:
                rr.font.bold = True
            elif j == 3:
                rr.font.color.rgb = GRAY
            if i % 2 == 0:
                set_cell_bg(c, "e8f5ee")

doc.add_paragraph()


# ── SECTION 5 — DEEP COMPARISON ──────────────────────────────
add_h1("5.  Head-to-Head Comparison")

compare = [
    ["Criterion", "Bridge Hub Code", "AI Functions"],
    ["Function count", "474  (73.8 %)", "168  (26.2 %)"],
    ["File count", "86 files", "33 files"],
    ["LLM providers", "none", "Claude Sonnet + GPT-4o + Gemini 2.5"],
    ["Response time", "instant (DB/cache)", "200 - 2 000 ms (LLM)"],
    ["Georgian language", "partial (UI labels)", "full bilingual prompts"],
    ["Learning ability", "static rules", "dynamic patterns + decay"],
    ["Conversation memory", "none", "10-turn session history"],
    ["DB usage", "all operations", "memory + patterns + cost logs"],
    ["Tests written", "48 unit tests", "8 dedicated AI tests"],
    ["CI/CD", "GitHub Actions -> Cloud Run", "API keys via GCP secrets"],
]

t5 = doc.add_table(rows=len(compare), cols=3)
t5.style = "Table Grid"
w5 = [Cm(4.5), Cm(6), Cm(6)]

for i, row_data in enumerate(compare):
    row = t5.rows[i]
    for j, txt in enumerate(row_data):
        c = row.cells[j]
        c.width = w5[j]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = pp.add_run(txt)
        rr.font.size = Pt(9)
        if i == 0:
            rr.font.bold = True
            rr.font.color.rgb = WHITE
            set_cell_bg(c, "1a1a2e")
        else:
            if j == 1:
                rr.font.color.rgb = HUB_COLOR
                if i % 2 == 0:
                    set_cell_bg(c, "eef2ff")
            elif j == 2:
                rr.font.color.rgb = AI_COLOR
                if i % 2 == 0:
                    set_cell_bg(c, "e8f5ee")

doc.add_paragraph()

# footer
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run("Bridge Hub v2  |  Statistics generated 2026-04-24  |  r.gelikoshvili@gmail.com")
rf.font.size = Pt(8)
rf.font.color.rgb = GRAY

output = r"C:\Users\Acer\Desktop\BridgeHub_Statistics.docx"
doc.save(output)
print("Saved:", output)
