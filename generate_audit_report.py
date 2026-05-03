"""Generate Bridge Hub Architecture Audit Report as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_RED   = RGBColor(0x8C, 0x3C, 0x2D)
BROWN      = RGBColor(0x4A, 0x37, 0x28)
CREAM_BG   = RGBColor(0xF3, 0xEC, 0xDC)
BLACK      = RGBColor(0x1A, 0x1A, 0x1A)
RED_WARN   = RGBColor(0xC0, 0x39, 0x2B)
GREEN_OK   = RGBColor(0x27, 0x76, 0x45)
ORANGE     = RGBColor(0xD3, 0x5F, 0x00)
BLUE_INFO  = RGBColor(0x1A, 0x53, 0x99)
GREY       = RGBColor(0x66, 0x66, 0x66)
TABLE_HDR  = RGBColor(0x8C, 0x3C, 0x2D)


def shade_cell(cell, hex_color: str):
    """Fill table cell background."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "start", "bottom", "end"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D4C5B0")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_RED
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "8C3C2D")
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BROWN


def normal(text, indent=0, color=None, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    run.font.color.rgb = color or BLACK
    return p


def bullet(text, level=0, color=None, bold=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.font.color.rgb = color or BLACK


def code_para(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = BROWN


def status_line(icon, severity, text):
    """[icon] [SEVERITY] text — coloured line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

    color_map = {
        "CRITICAL": RED_WARN,
        "HIGH":     ORANGE,
        "MEDIUM":   BLUE_INFO,
        "LOW":      GREY,
        "OK":       GREEN_OK,
        "WARNING":  ORANGE,
    }
    c = color_map.get(severity, BLACK)

    r1 = p.add_run(f"{icon}  ")
    r1.font.size = Pt(11)
    r1.font.color.rgb = c

    r2 = p.add_run(f"[{severity}]  ")
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = c

    r3 = p.add_run(text)
    r3.font.size = Pt(10.5)
    r3.font.color.rgb = BLACK


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, "8C3C2D")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, row in enumerate(rows):
        drow = t.rows[ri + 1]
        bg = "FFF8F0" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            shade_cell(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table


# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
r = p.add_run("Bridge Hub")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = DARK_RED

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Full Architecture Audit & Competitive Analysis Report")
r2.font.size = Pt(16)
r2.font.color.rgb = BROWN

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(f"შედგენის თარიღი: 2026-05-03  |  Revision: 00191-d8l  |  Region: europe-west1")
r3.font.size = Pt(10)
r3.font.color.rgb = GREY
r3.italic = True

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

heading1("Executive Summary")
normal(
    "Bridge Hub არის AI-პირველი, ქართული აუდიტორული და ბუღალტრული SaaS პლატფორმა, "
    "რომელიც Google Cloud Run-ზე მუშაობს (europe-west1). სისტემა შეიცავს 287 Python "
    "ფაილს, 47,734 კოდის სტრიქონს, 92 სერვერ-მხარეს ცხრილს PostgreSQL-ში და "
    "227 REST API endpoint-ს. ეს რეპორტი მოიცავს სრულ კოდბაზის ანალიზს, "
    "სუსტ ადგილებს, მონაცემთა ბაზის სქემას, performance-ს და შედარებას "
    "Xero/QuickBooks-თან.",
    color=BROWN
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — CODEBASE METRICS
# ═══════════════════════════════════════════════════════════════════════════════

heading1("1. Codebase მეტრიკა")

heading2("1.1 ზომა და სტრუქტურა")

add_table(
    ["მეტრიკა", "მნიშვნელობა", "კომენტარი"],
    [
        ["Python ფაილები (app/)", "287", "tests/ გარეშე"],
        ["კოდის სტრიქონები (LOC)", "47,734", "კომენტარები ჩათვლით"],
        ["Route ფაილები", "85", "app/api/routes_*.py"],
        ["@router დეკორატორები", "413", "ყველა HTTP method"],
        ["უნიკალური endpoints", "227", "GET:130 POST:86 DELETE:7 PUT:3 PATCH:1"],
        ["Services (app/api/services/)", "70 ფაილი", ""],
        ["HTML გვერდები (static/)", "33", "approval, documents, inventory..."],
        ["Test ფაილები", "70", "unit + integration"],
        ["Test ფუნქციები", "300", "137 unit, 163 integration"],
        ["DB ცხრილები (production)", "92", "PostgreSQL bridgehub DB"],
        ["Background async loops", "3", "inbox poller, pattern decay, warm-up"],
        ["Archived/dead route files", "18", "app/api/archive/"],
    ],
    col_widths=[6, 4, 7]
)

heading2("1.2 ყველაზე დიდი ფაილები")
add_table(
    ["ფაილი", "LOC", "შეფასება"],
    [
        ["app/api/routes_documents.py", "1,433", "⚠️ სახელება — split გაიზიმება"],
        ["app/api/services/accounting_rules.py", "1,142", "✅ domain logic — ok"],
        ["app/api/services/approval_service.py", "858", "✅ core — ok"],
        ["app/api/services/ai_chat_service.py", "815", "✅ AI logic — ok"],
        ["app/api/routes_outgoing.py", "776", "⚠️ PDF + routes — split გაიზიმება"],
        ["app/api/services/posting_service.py", "722", "✅ core — ok"],
        ["app/api/transaction_classifier.py", "720", "✅ AI engine — ok"],
        ["app/api/engines/pattern_engine.py", "679", "✅ Immutable Core"],
        ["app/api/routes_auth.py", "657", "⚠️ ძალიან დიდი auth ფაილი"],
        ["app/api/routes_claude_chat.py", "601", "✅ ok"],
    ],
    col_widths=[8, 2.5, 6.5]
)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — ARCHITECTURE WEAKNESSES
# ═══════════════════════════════════════════════════════════════════════════════

heading1("2. არქიტექტურული სისუსტეები")

heading2("2.1 Double-Entry Validation (Dr = Cr)")
status_line("✅", "OK", "IMPLEMENTED — posting_service.py:131")
normal(
    "posting_service.py-ში არსებობს სრული Dr=Cr შემოწმება. "
    "თუ debit_total != credit_total, სისტემა hard error-ს აბრუნებს "
    "და posting-ს ბლოკავს. ეს არის ბუღალტრული მთავარი ვალიდაცია.",
    indent=1
)
code_para("if round(debit_total, 2) != round(credit_total, 2):")
code_para('    return f"დებეტი და კრედიტი არ ემთხვევა (Dr={debit_total}, Cr={credit_total})"')

doc.add_paragraph()
heading2("2.2 Idempotency (ორმაგი posting-ის პრევენცია)")
status_line("⚠️", "WARNING", "PARTIAL — duplicate warning მხოლოდ, hard block არ არის")
normal(
    "idempotency_keys ცხრილი არსებობს DB-ში და index-ებიც დადებულია "
    "(idx_idempotency_keys_lookup), მაგრამ POST /posting/apply endpoint-ზე "
    "ის არ გამოიყენება. posting_service.py:531-ზე არსებობს "
    "duplicate warning (±1 GEL, ±3 დღე), მაგრამ ეს მხოლოდ გაფრთხილებაა — "
    "ვერ ბლოკავს ორ identically-ს request-ს milliseconds სხვაობით.",
    indent=1
)
normal("გამოსწორება: POST /posting/apply-ზე idempotency_keys-ის გამოყენება.", indent=1, color=ORANGE, bold=True)

doc.add_paragraph()
heading2("2.3 Tenant Isolation (cross-tenant data leak)")
status_line("❌", "HIGH", "16 route ფაილში SELECT..WHERE-ს tenant_id არ აქვს")
normal(
    "Multi-tenant სისტემაში ყველა SELECT query-ს WHERE tenant_id = %s "
    "უნდა ჰქონდეს. შემდეგ ფაილებში ეს ფილტრი გამოტოვებულია:",
    indent=1
)

add_table(
    ["Route ფაილი", "query-ების რაოდენობა", "რისკის დონე"],
    [
        ["routes_coa.py",              "6",  "⚠️ MEDIUM — Chart of Accounts შეიძლება global-ი იყოს"],
        ["routes_search.py",           "4",  "🔴 HIGH — full-text search cross-tenant"],
        ["routes_security.py",         "5",  "🔴 HIGH — security events cross-tenant"],
        ["routes_expenses.py",         "3",  "🔴 HIGH — expense data cross-tenant"],
        ["routes_invoices.py",         "3",  "🔴 HIGH — invoice cross-tenant"],
        ["routes_contracts.py",        "2",  "🔴 HIGH"],
        ["routes_cost_center.py",      "2",  "🔴 HIGH"],
        ["routes_employee_portal.py",  "2",  "🔴 HIGH"],
        ["routes_audit_engine.py",     "1",  "⚠️ MEDIUM"],
        ["routes_audit_log.py",        "1",  "⚠️ MEDIUM"],
        ["routes_audit_trail.py",      "1",  "⚠️ MEDIUM"],
        ["routes_balance_ge.py",       "1",  "⚠️ MEDIUM"],
        ["routes_budget.py",           "1",  "⚠️ MEDIUM"],
        ["routes_crm.py",              "1",  "⚠️ MEDIUM"],
        ["routes_rbac.py",             "1",  "⚠️ LOW — RBAC table შეიძლება global-ი იყოს"],
        ["routes_webhooks_v2.py",      "1",  "⚠️ MEDIUM"],
    ],
    col_widths=[6.5, 4, 6.5]
)

doc.add_paragraph()
heading2("2.4 Async vs Sync თანაფარდობა")
status_line("⚠️", "HIGH", "13% async (69 async def / 452 sync def)")
normal(
    "FastAPI async event loop-ზე მუშაობს, მაგრამ psycopg2 sync driver-ია. "
    "452 sync route handler threadpool-ში გადადის DB call-ის დროს. "
    "ეს პრობლემა არ გამოიჩენება მცირე load-ზე, მაგრამ high concurrency-ზე "
    "(100+ req/sec) threadpool exhaustion-ი შეიძლება მოხდეს.",
    indent=1
)
normal("გამოსწორება: asyncpg + SQLAlchemy async — throughput 3-5x გაიზარდება.", indent=1, color=ORANGE, bold=True)

doc.add_paragraph()
heading2("2.5 Error Handling")
status_line("⚠️", "MEDIUM", "except Exception — მაღალი სიხშირე კრიტიკულ ფაილებში")
normal("bare except: — 0 (✅ კარგია, ყველა typed-ია)", indent=1, color=GREEN_OK)
normal("except Exception count კრიტიკულ ფაილებში:", indent=1)
add_table(
    ["ფაილი", "except Exception count"],
    [
        ["routes_documents.py",   "16"],
        ["routes_claude_chat.py", "16"],
        ["routes_outgoing.py",    "14"],
        ["routes_inventory.py",   "13"],
        ["routes_approval.py",    "12"],
        ["routes_search.py",       "8"],
        ["routes_learning.py",     "8"],
        ["routes_integrations.py", "8"],
    ],
    col_widths=[10, 4]
)
normal(
    "პრობლემა: except Exception ბლოკები ხშირად ერთნაირ generic error response-ს "
    "აბრუნებს და კონკრეტულ exception type-ებს არ ასხვავებს. "
    "ეს debugging-ს ართულებს.",
    indent=1, color=GREY
)

doc.add_paragraph()
heading2("2.6 Rate Limiting")
status_line("⚠️", "MEDIUM", "11/85 route ფაილზე @limiter.limit — დანარჩენი unprotected")
normal("SlowAPIMiddleware გლობალურად active-ია ✅", indent=1, color=GREEN_OK)
normal("@limiter.limit დეკორატორები:", indent=1)
add_table(
    ["Route ფაილი", "decorated endpoints"],
    [
        ["routes_auth.py",        "8  (login: 5/min ✅)"],
        ["routes_outgoing.py",    "6"],
        ["routes_reports.py",     "5"],
        ["routes_posting.py",     "5"],
        ["routes_documents.py",   "5"],
        ["routes_approval.py",    "5"],
        ["routes_export_v2.py",   "4"],
        ["routes_recurring.py",   "3"],
        ["routes_ocr.py",         "2"],
        ["routes_bank_csv.py",    "1"],
        ["routes_aging.py",       "1"],
    ],
    col_widths=[8, 6]
)
normal("დაუცველი: routes_search, routes_crm, routes_dashboard_live, routes_invoices", indent=1, color=RED_WARN)

doc.add_paragraph()
heading2("2.7 Connection Pool")
status_line("⚠️", "MEDIUM", "DB_POOL_MAX=4 — undersized for containerConcurrency=100")
normal(
    "Cloud Run instance-ი 100 simultaneous request-ს ატარებს (containerConcurrency=100), "
    "მაგრამ DB connection pool მაქსიმუმ 4 კავშირს ხსნის. "
    "მაღალი load-ის დროს 96 request ელოდება თავისუფალ კავშირს.",
    indent=1
)
add_table(
    ["პარამეტრი", "მიმდინარე", "რეკომენდაცია"],
    [
        ["DB_POOL_MIN", "1",   "2"],
        ["DB_POOL_MAX", "4",   "16-20"],
        ["containerConcurrency", "100", "80 (DB pool-თან შესაბამისი)"],
    ],
    col_widths=[5, 4, 8]
)

doc.add_paragraph()
heading2("2.8 Dead Code (archived routes)")
status_line("ℹ️", "LOW", "18 ფაილი app/api/archive/ — unused, კოდბაზის noise")
normal("archived ფაილები:", indent=1)
archived = [
    "routes_bank.py", "routes_dashboard_full.py", "routes_dashboard_mobile.py",
    "routes_dashboard_ui.py", "routes_dashboard_v2.py", "routes_finance.py",
    "routes_finance_engine.py", "routes_financial_statements.py", "routes_firestore.py",
    "routes_fpa.py", "routes_gates.py", "routes_launch.py", "routes_pipeline.py",
    "routes_reconciliation_v2.py", "routes_reports_dashboard.py", "routes_strategy.py",
    "routes_supervisor.py"
]
for f in archived:
    bullet(f, color=GREY)

doc.add_paragraph()
heading2("2.9 Security Headers")
status_line("✅", "OK", "სრული security header ნაკრები დადგენილია")
add_table(
    ["Header", "მნიშვნელობა"],
    [
        ["X-Content-Type-Options", "nosniff"],
        ["X-Frame-Options", "SAMEORIGIN"],
        ["X-XSS-Protection", "1; mode=block"],
        ["Strict-Transport-Security", "max-age=63072000; includeSubDomains; preload"],
        ["Referrer-Policy", "strict-origin-when-cross-origin"],
        ["Permissions-Policy", "camera=(), microphone=(), geolocation=()"],
        ["Content-Security-Policy", "default-src 'self' + CDN whitelist (jsdelivr, unpkg, googleapis)"],
    ],
    col_widths=[6, 11]
)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════════════════════════

heading1("3. მონაცემთა ბაზის სქემა")

heading2("3.1 ცხრილების სია (92 ცხრილი)")
add_table(
    ["ცხრილი", "სვეტები", "ჩანაწერები", "კომენტარი"],
    [
        ["audit_log",           "17", "9,434",  "ყველაზე დიდი — 9K audit event"],
        ["audit_events",         "6", "402",    ""],
        ["bank_transactions",   "10", "284",    ""],
        ["learning_feedback",   "25", "57",     "AI learning data"],
        ["llm_cost_log",         "7", "73",     "LLM API ხარჯების log"],
        ["posting_logs",        "10", "53",     ""],
        ["learning_patterns",   "22", "33",     ""],
        ["outgoing_invoices",   "35", "103",    "ყველაზე სრული schema"],
        ["journal_drafts",      "63", "8",      "ყველაზე კომპლექსური — 63 სვეტი"],
        ["tenants",             "25", "5",      "multi-tenant core"],
        ["users",                "8", "14",     ""],
        ["waybills",            "24", "7",      ""],
        ["tax_invoices",        "21", "11",     ""],
        ["erp_posting_memory",  "22", "20",     "AI posting memory"],
        ["learning_deltas",     "10", "16",     ""],
        ["human_reviews",       "11", "16",     "manual review queue"],
        ["exchange_rates",       "4", "12",     "FX rates"],
        ["transaction_memory",  "11", "22",     ""],
        ["contracts",           "16", "7",      ""],
        ["invoices",            "16", "4",      ""],
        ["expenses",            "14", "4",      ""],
        ["customers",           "12", "3",      ""],
        ["idempotency_keys",     "6", "0",      "⚠️ table ready, wiring incomplete"],
        ["triangle_matches",    "15", "0",      "waybill+tax_invoice+po matching"],
        ["vat_returns",          "9", "0",      "VAT declaration records"],
        ["period_locks",         "8", "0",      "accounting period locking"],
    ],
    col_widths=[5.5, 2.5, 3.5, 5.5]
)
normal("სხვა 66 ცხრილი: employees, purchase_orders, warehouses, cost_centers, webhooks, "
       "pension_transfers, async_queue, rate_limits, search_index და სხვ.", color=GREY)

doc.add_paragraph()
heading2("3.2 ყველაზე დიდი ცხრილები (disk size)")
add_table(
    ["ცხრილი", "ზომა", "შენიშვნა"],
    [
        ["audit_log",           "2,160 kB", "სწრაფად იზრდება — rotation სჭირდება"],
        ["tenants",             "1,544 kB", "⚠️ 5 ჩანაწერზე ძალიან დიდია — JSONB bloat"],
        ["email_documents",       "768 kB", "3 ჩანაწერი — large BLOB content-ი"],
        ["processed_documents",   "488 kB", "OCR result-ების JSONB"],
        ["journal_drafts",        "256 kB", "63-column JSONB heavy schema"],
        ["bank_transactions",     "248 kB", "284 rows"],
        ["outgoing_invoices",     "152 kB", "103 rows"],
        ["learning_patterns",     "152 kB", "33 rows — rich pattern data"],
    ],
    col_widths=[5, 3, 9]
)
normal("⚠️ tenants ცხრილი 5 ჩანაწერზე 1.5MB-ია — JSONB სვეტებში ჩვეულება bloat. "
       "VACUUM FULL გაიზიმება.", indent=1, color=ORANGE)

doc.add_paragraph()
heading2("3.3 Index-ების მიმოხილვა (196 index)")
normal("კარგად index-ებული ცხრილები:", bold=True)
add_table(
    ["ცხრილი", "index-ების რ-ბა", "ყველაზე მნიშვნელოვანი"],
    [
        ["journal_drafts",    "12", "tenant_status_created, fingerprint, confidence"],
        ["outgoing_invoices",  "5", "tenant, status, buyer_inn, created"],
        ["tax_invoices",       "6", "tenant, buyer_inn, seller_inn, date"],
        ["waybills",           "6", "tenant, buyer_inn, seller_inn, date"],
        ["audit_log",          "5", "action, event_time, resource"],
        ["learning_patterns",  "5", "tenant_type, usage_count, last_used"],
        ["idempotency_keys",   "3", "UNIQUE composite key ✅"],
    ],
    col_widths=[5, 3, 9]
)
normal("შესაძლო missing index-ები:", bold=True)
add_table(
    ["ცხრილი", "სვეტი", "რისკი"],
    [
        ["expenses",          "tenant_id", "MEDIUM — full scan on expense queries"],
        ["contracts",         "tenant_id", "MEDIUM"],
        ["customers",         "tenant_id", "MEDIUM"],
        ["bank_transactions", "account_id", "LOW — only 284 rows now"],
        ["invoices",          "tenant_id", "MEDIUM"],
    ],
    col_widths=[4.5, 4, 8.5]
)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — PERFORMANCE
# ═══════════════════════════════════════════════════════════════════════════════

heading1("4. Performance ანალიზი")

heading2("4.1 Production Response Times")
add_table(
    ["Endpoint", "HTTP Status", "Response Time", "შეფასება"],
    [
        ["GET /health/",                "200", "4.95s", "⚠️ cold start / DB connect warmup"],
        ["GET /static/approval.html",   "200", "0.60s", "✅ static file serving — ok"],
    ],
    col_widths=[6, 2.5, 3, 5.5]
)
normal(
    "4.95s /health response-ი კი cold start-ის ნიშანია (DB connection pool-ის გახსნა). "
    "min_instances=1-ის გამო cold start-ი იშვიათია, მაგრამ startup-cpu-boost=true "
    "ჩართულია ✅, რაც startup-ს ამცირებს.",
    indent=1, color=GREY
)

doc.add_paragraph()
heading2("4.2 Cloud Run კონფიგურაცია")
add_table(
    ["პარამეტრი", "მიმდინარე მნიშვნელობა", "რეკომენდაცია"],
    [
        ["min_instances",       "1",   "1 ✅"],
        ["max_instances",       "4",   "6-8 (traffic growth-ისთვის)"],
        ["containerConcurrency","100", "80 (DB pool-თან sync-ში)"],
        ["timeout",             "300s","300s ✅ (OCR-ისთვის საჭიროა)"],
        ["startup-cpu-boost",   "true","true ✅"],
        ["DB_POOL_MIN",         "1",   "2"],
        ["DB_POOL_MAX",         "4",   "16"],
    ],
    col_widths=[5, 4, 8]
)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — TEST COVERAGE
# ═══════════════════════════════════════════════════════════════════════════════

heading1("5. Test Coverage ანალიზი")

heading2("5.1 ზოგადი სტატისტიკა")
add_table(
    ["მეტრიკა", "მნიშვნელობა"],
    [
        ["Unit tests (collected)", "137"],
        ["Integration test functions", "163"],
        ["სულ test functions", "300"],
        ["Route ფაილები dedicated test-ით", "7 / 85  (8%)"],
        ["Test ფაილები სულ", "70"],
    ],
    col_widths=[8, 9]
)

doc.add_paragraph()
heading2("5.2 კარგად დაფარული არეები")
good_tested = [
    "Approval pipeline (approve/reject/correct/batch)",
    "Document intelligence (INN extraction, OCR, confidence scoring)",
    "Upload workflow (FOREIGN confidence, duplicate detection)",
    "Georgia tax pack (VAT, payroll PIT/SSIC, CIT)",
    "Posting integrity (Dr=Cr validation)",
    "Learning engine (pattern learning, decay)",
    "Tenant isolation (cross-tenant data leak tests)",
    "RBAC (permission enforcement)",
    "Speed optimization (response time benchmarks)",
    "Approval race condition (concurrent request handling)",
]
for t in good_tested:
    bullet(t, color=GREEN_OK)

doc.add_paragraph()
heading2("5.3 0 Test-ის მქონე კრიტიკული route-ები")
untested_critical = [
    ("routes_auth.py",       "JWT login, registration, password reset — კრიტიკული"),
    ("routes_payroll.py",    "PIT/SSIC/pension გამოთვლები"),
    ("routes_outgoing.py",   "invoice lifecycle, PDF generation, multi-currency"),
    ("routes_recurring.py",  "recurring template + auto-send"),
    ("routes_aging.py",      "overdue reminders, SMTP"),
    ("routes_fx.py",         "exchange rate sync"),
    ("routes_fixed_assets.py","depreciation calculation"),
    ("routes_closing.py",    "period closing logic"),
    ("routes_2fa.py",        "TOTP authentication"),
    ("routes_bank_sync.py",  "bank transaction sync"),
]
add_table(
    ["Route ფაილი", "შეფასება"],
    untested_critical,
    col_widths=[6, 11]
)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — SECURITY AUDIT
# ═══════════════════════════════════════════════════════════════════════════════

heading1("6. Security Audit")

heading2("6.1 Authentication & Authorization")
sec_features = [
    ("✅", "JWT (PyJWT) — stateless auth, BridgeHub2026JWT secret"),
    ("✅", "TOTP / 2FA (pyotp + qrcode) — /api/2fa/ routes"),
    ("✅", "OAuth2 — /api/oauth/ routes"),
    ("✅", "API Key auth — api_keys ცხრილი + SHA-256 hashing"),
    ("✅", "bcrypt password hashing (v4.2.1)"),
    ("✅", "blocked_ips ცხრილი + auto-block logic"),
    ("✅", "security_events audit trail"),
    ("✅", "require_permission() RBAC on approval + document endpoints"),
    ("✅", "Rate limiting: /auth/login 5/minute"),
    ("⚠️", "idempotency_keys table ready — wiring to /posting/apply missing"),
    ("⚠️", "JWT secret hardcoded in env var — Key Management Service-ი სჯობს"),
]
for icon, text in sec_features:
    color = GREEN_OK if icon == "✅" else ORANGE
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{icon}  ")
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size  = Pt(10.5)
    r2.font.color.rgb = color

doc.add_paragraph()
heading2("6.2 Dependencies Security")
add_table(
    ["Package", "მიმდინარე", "Latest", "შენიშვნა"],
    [
        ["bcrypt",       "4.2.1",  "5.0.0", "⚠️ v5 — breaking API changes"],
        ["cryptography", "46.0.5", "47.0.0","⚠️ security patch შეიძლება"],
        ["fastapi",      "0.135.2","0.136.1","minor update"],
        ["grpcio",       "1.78.0", "1.80.0", "minor update"],
        ["httpx",        "0.27.2", "0.28.1", "minor update"],
    ],
    col_widths=[4, 3, 3, 7]
)
normal("pip-audit vulnerability scan: კრიტიკული CVE-ები ვერ აღმოჩნდა.", color=GREEN_OK)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — COMPETITIVE ANALYSIS (XERO / QUICKBOOKS / 1C)
# ═══════════════════════════════════════════════════════════════════════════════

heading1("7. კონკურენტული ანალიზი — Xero / QuickBooks / 1C / Balance.ge")

heading2("7.1 Bridge Hub > კონკურენტები (უპირატესობები)")
better = [
    "AI auto-classification with learning engine — Xero/QB-ს manual rules აქვს, Bridge Hub-ს pattern learning",
    "Georgian tax compliance built-in: VAT 18%, payroll PIT (20%) + SSIC, CIT (15%), RS.ge integration, waybill triangle matching",
    "Document OCR pipeline: PDF/image → journal draft, ბუღალტრული ჩანაწერი ავტომატიზირებულია",
    "LLM chat with full accounting context (Claude + OpenAI) — არც Xero, არც QB-ს ეს არ აქვს",
    "Georgian-language UI + bilingual error messages (Georgian + English)",
    "Confidence scoring on auto-generated drafts — transparency of AI decisions",
    "Pattern learning from human corrections — improves over time",
    "Waybill + Tax Invoice triangle matching (RS.ge specific requirement)",
    "Recurring invoice templates with auto-send",
    "Aging reminders with Georgian HTML email templates",
]
for b in better:
    bullet(b, color=GREEN_OK)

doc.add_paragraph()
heading2("7.2 Bridge Hub ≈ კონკურენტები (ტოლი ფუნქციონალი)")
equal = [
    "Double-entry enforcement — ყველა აკეთებს; Bridge Hub Dr=Cr hard validation ✅",
    "Multi-currency — Bridge Hub: 7 currencies (GEL/USD/EUR/GBP/TRY/RUB/CHF); Xero: 160+",
    "Invoice lifecycle — draft→finalized; similar to Xero",
    "Audit trail — present; Xero-ს მეტი granularity აქვს",
    "API-first design — ყველა სამი",
    "Bank CSV import — basic; Xero-ს more polished UI",
    "Multi-tenant — Bridge Hub SaaS, 1C local",
]
for e in equal:
    bullet(e, color=BLUE_INFO)

doc.add_paragraph()
heading2("7.3 Bridge Hub < კონკურენტები (სუსტი ადგილები vs Xero/QB)")
missing = [
    ("Mobile app",
     "Xero/QB-ს iOS/Android app აქვს; Bridge Hub — web only"),
    ("Financial Statements UI",
     "routes_financial_statements.py არსებობს, მაგრამ rendered P&L/Balance Sheet UI არ არის"),
    ("Budget vs Actual report",
     "budgets ცხრილი არსებობს, variance report — არ არის"),
    ("Inventory valuation report",
     "FIFO/LIFO/avg tracking მუშაობს, report page — არ არის"),
    ("Payroll payslip PDF",
     "payroll calculation მუშაობს, printable payslip — არ არის"),
    ("Live bank feed",
     "CSV import only; Xero-ს direct bank API (Open Banking) connection აქვს"),
    ("Quotes / Estimates",
     "Invoice-ამდე quote stage — არ არის"),
    ("Project/job costing",
     "cost_centers არსებობს, project profitability report — არ არის"),
    ("Recurring expense tracking",
     "Recurring invoices გაქვს, recurring expenses — არ გაქვს"),
    ("Client portal (pay online)",
     "routes_client_portal.py არსებობს; online payment integration — არ არის"),
]
add_table(
    ["რაც აკლია", "კომენტარი"],
    missing,
    col_widths=[5, 12]
)

doc.add_paragraph()
heading2("7.4 Balance.ge / RS.ge შედარება (ქართული კონკურენტები)")
add_table(
    ["ფუნქცია", "Bridge Hub", "Balance.ge", "RS.ge portal"],
    [
        ["AI journal auto-fill",    "✅ Claude/OpenAI",   "❌",  "❌"],
        ["OCR document pipeline",   "✅ Vision API",      "⚠️ basic", "❌"],
        ["Learning engine",         "✅ pattern decay",   "❌",  "❌"],
        ["Waybill management",      "✅ full",            "✅",  "✅"],
        ["Tax invoice",             "✅ full",            "✅",  "✅"],
        ["VAT declaration",         "✅ routes_tax.py",   "✅",  "✅ native"],
        ["Payroll (PIT/SSIC)",      "✅ georgia_pack.py", "✅",  "❌"],
        ["Multi-currency",          "✅ 7 currencies",    "⚠️",  "❌"],
        ["REST API",                "✅ 227 endpoints",   "⚠️ partial", "❌"],
        ["Open source / SaaS",      "✅ SaaS",            "SaaS","gov portal"],
    ],
    col_widths=[5.5, 4, 3.5, 4]
)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — PRIORITY ACTION PLAN
# ═══════════════════════════════════════════════════════════════════════════════

heading1("8. პრიორიტეტული სამოქმედო გეგმა")

heading2("8.1 CRITICAL — Production Risk (დაუყოვნებლივ)")
critical_actions = [
    ("DB_POOL_MAX 4 → 16",
     "Connection starvation under load. Cloud Run concurrency=100 vs pool=4. "
     "ერთი env var ცვლილება — gcloud run services update."),
    ("Tenant isolation fix: routes_search.py, routes_expenses.py, routes_invoices.py",
     "Cross-tenant data leakage. 3 ფაილში WHERE clause-ზე tenant_id = %s დამატება."),
]
for title, desc in critical_actions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f"🔴  {title}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RED_WARN
    normal(desc, indent=2, color=BROWN)

doc.add_paragraph()
heading2("8.2 HIGH — Quality (მომდევნო sprint)")
high_actions = [
    ("POST /posting/apply — idempotency_keys wire-up",
     "idempotency_keys table + index ready. Endpoint-ზე ON CONFLICT DO NOTHING დამატება."),
    ("routes_auth.py tests",
     "JWT, password reset, 2FA — 0 unit tests. კრიტიკული path სრულად untested."),
    ("routes_search.py @limiter.limit",
     "Full-text search without rate limit — scraping/DDoS vector."),
    ("routes_outgoing.py + routes_recurring.py + routes_aging.py tests",
     "3 ახალი feature — 0 tests. invoice lifecycle, email sending."),
]
for title, desc in high_actions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f"🟠  {title}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ORANGE
    normal(desc, indent=2, color=BROWN)

doc.add_paragraph()
heading2("8.3 MEDIUM — Feature Gaps (შემდეგი quarter)")
medium_actions = [
    "Financial Statements HTML page — P&L + Balance Sheet rendered (routes_financial_statements.py დაასრულება)",
    "Budget vs Actual variance report — budgets ცხრილი ready",
    "Payroll payslip PDF — reportlab template (georgia_pack.py payroll function + PDF)",
    "audit_log rotation — 9K+ rows, სწრაფად იზრდება; partition by month ან pg_partman",
    "asyncpg migration — sync→async DB driver, throughput 3-5x",
    "tenants ცხრილი VACUUM FULL — 5 rows, 1.5MB JSONB bloat",
]
for a in medium_actions:
    bullet(f"🔵  {a}", color=BLUE_INFO)

doc.add_paragraph()
heading2("8.4 LOW — Housekeeping")
low_actions = [
    "app/api/archive/ — 18 dead files წაშლა (confirm before delete)",
    "bcrypt 4.2 → 5.0 upgrade (API breaking — checkpw signature changed)",
    "DB_POOL_MIN 1 → 2 (faster first-request warmup)",
    "routes_documents.py split — 1,433 LOC → document_upload + document_management",
]
for a in low_actions:
    bullet(f"⚪  {a}", color=GREY)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — SUMMARY SCORECARD
# ═══════════════════════════════════════════════════════════════════════════════

heading1("9. შეჯამება — Scorecard")

add_table(
    ["კატეგორია", "ქულა", "შეფასება"],
    [
        ["Double-entry enforcement",   "10/10", "✅ hard Dr=Cr block"],
        ["Security headers",           "10/10", "✅ სრული OWASP header set"],
        ["Authentication",             " 9/10", "✅ JWT+2FA+OAuth; -1 for hardcoded secret"],
        ["DB indexing",                " 8/10", "✅ კარგი coverage; 5 missing indexes"],
        ["Georgian tax compliance",    "10/10", "✅ VAT+payroll+CIT+RS.ge"],
        ["AI / automation",            " 9/10", "✅ OCR+LLM+learning; -1 no mobile"],
        ["Tenant isolation",           " 6/10", "⚠️ 16 routes with gaps"],
        ["Test coverage",              " 4/10", "❌ 8% route coverage"],
        ["Connection pool sizing",     " 3/10", "❌ max=4 vs concurrency=100"],
        ["Async performance",          " 5/10", "⚠️ 13% async, sync psycopg2"],
        ["Rate limiting",              " 6/10", "⚠️ global middleware + 11 files"],
        ["Financial reporting UI",     " 5/10", "⚠️ P&L/Balance Sheet missing"],
        [" ᲡᲣᲚ",                       "85/120", "71% — Production-ready core"],
    ],
    col_widths=[7, 2.5, 7.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Bridge Hub — production-ready AI accounting platform. "
    "3 critical fix-ის შემდეგ (pool size, tenant isolation, idempotency) "
    "სისტემა enterprise-grade-ია."
)
r.font.size  = Pt(11)
r.font.color.rgb = DARK_RED
r.bold = True
r.italic = True


# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Bridge Hub Architecture Audit Report  |  Generated 2026-05-03  |  Confidential")
r.font.size  = Pt(9)
r.font.color.rgb = GREY
r.italic = True


# ── Save ──────────────────────────────────────────────────────────────────────
out = "Bridge_Hub_Architecture_Audit_2026.docx"
doc.save(out)
print(f"Saved: {out}")
